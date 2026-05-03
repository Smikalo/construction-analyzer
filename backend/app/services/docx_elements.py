"""Helpers for normalizing DOCX evidence into typed document elements."""

from __future__ import annotations

from dataclasses import replace
from datetime import datetime
from typing import Any

import docx
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.services.document_elements import DocumentElement

from . import table_elements

DOCX_ELEMENT_TYPE = "file_summary"
DOCX_EXTRACTION_MODE = "docx_summary"
DOCX_SUBJECT = "engineering_narrative"


def extract_docx(
    path: str,
    *,
    source: str,
    document_id: str | None = None,
) -> list[DocumentElement]:
    """Extract a deterministic summary plus ordered body elements from DOCX."""
    document = docx.Document(path)
    core_properties = document.core_properties

    title = _coerce_text(core_properties.title)
    author = _coerce_text(core_properties.author)
    created = _format_created(core_properties.created)
    paragraph_count = len(document.paragraphs)

    metadata: dict[str, Any] = {
        "subject": DOCX_SUBJECT,
        "paragraph_count": paragraph_count,
    }
    if title:
        metadata["docx_title"] = title
    if author:
        metadata["docx_author"] = author
    if created:
        metadata["docx_created"] = created

    elements = [
        DocumentElement(
            document_id=document_id,
            source=source,
            path=path,
            page=None,
            element_type=DOCX_ELEMENT_TYPE,
            extraction_mode=DOCX_EXTRACTION_MODE,
            content="\n".join(
                [
                    f"Title: {title}",
                    f"Author: {author}",
                    f"Created: {created}",
                    f"Paragraphs: {paragraph_count}",
                ]
            ),
            confidence=None,
            warnings=(),
            metadata=metadata,
        )
    ]

    block_index = 0
    current_section_heading: str | None = None
    body = document.element.body

    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            paragraph = Paragraph(child, parent=document._body)
            text = paragraph.text.strip()
            if not text:
                continue

            style_name = paragraph.style.name if paragraph.style else ""
            element_type = "heading" if style_name.startswith("Heading") else "paragraph"
            element_metadata: dict[str, Any] = {
                "block_index": block_index,
                "style_name": style_name,
                "subject": DOCX_SUBJECT,
            }
            _add_section_heading(element_metadata, current_section_heading)

            elements.append(
                DocumentElement(
                    document_id=document_id,
                    source=source,
                    path=path,
                    page=None,
                    element_type=element_type,
                    extraction_mode=(
                        "docx_heading" if element_type == "heading" else "docx_paragraph"
                    ),
                    content=text,
                    confidence=None,
                    warnings=(),
                    metadata=element_metadata,
                )
            )
            block_index += 1
            if element_type == "heading":
                current_section_heading = text
            continue

        if child.tag == qn("w:tbl"):
            table = Table(child, parent=document._body)
            rows = [[cell.text for cell in row.cells] for row in table.rows]
            extra_warnings = (
                ("unsupported_structure",)
                if any(cell.tables for row in table.rows for cell in row.cells)
                else ()
            )
            table_metadata = _table_metadata(block_index, current_section_heading)
            table_element = table_elements.table_element_from_rows(
                rows,
                document_id=document_id,
                source=source,
                path=path,
                page=None,
                warnings=extra_warnings,
                metadata=table_metadata,
            )
            if table_element is None:
                continue
            elements.append(replace(table_element, extraction_mode="docx_table"))
            block_index += 1

    return elements


def _add_section_heading(metadata: dict[str, Any], section_heading: str | None) -> None:
    if section_heading is not None:
        metadata["section_heading"] = section_heading


def _table_metadata(
    block_index: int,
    current_section_heading: str | None,
) -> dict[str, Any]:
    metadata: dict[str, Any] = {
        "block_index": block_index,
        "subject": DOCX_SUBJECT,
    }
    _add_section_heading(metadata, current_section_heading)
    return metadata


def _coerce_text(value: object | None) -> str:
    if value is None:
        return ""
    return str(value)


def _format_created(value: datetime | None) -> str:
    if value is None:
        return ""
    return value.isoformat()


__all__ = ["extract_docx"]
