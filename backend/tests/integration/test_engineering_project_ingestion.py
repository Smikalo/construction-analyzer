"""Integration proof for mixed engineering folder ingestion and recall provenance."""

from __future__ import annotations

import hashlib
import io
from dataclasses import dataclass
from pathlib import Path

import docx
import pytest
from fastapi.testclient import TestClient
from openpyxl import Workbook

import app.services.converted_drawing_elements as converted_drawing_elements
from app.agent.tools import build_kb_tools
from app.kb.fake import FakeKB
from app.services.document_registry import lifespan_document_registry
from app.services.engineering_converters import ConversionResult
from app.services.ingestion import ingest_directory

pytestmark = [pytest.mark.integration]


@dataclass(frozen=True, slots=True)
class FakePage:
    text: str

    def extract_text(self) -> str | None:
        return self.text


class RecordingEngineeringConverter:
    def __init__(self, source_path: Path, converted_path: Path) -> None:
        self._source_path = source_path
        self._converted_path = converted_path
        self.calls: list[str] = []

    def convert(self, source_path: str) -> ConversionResult:
        self.calls.append(source_path)
        assert source_path == str(self._source_path)

        self._converted_path.parent.mkdir(parents=True, exist_ok=True)
        self._converted_path.write_bytes(b"%PDF-1.7\n")
        return ConversionResult(
            success=True,
            status="success",
            output_path=str(self._converted_path),
            warnings=("converter_note",),
            error=None,
            diagnostics={
                "layers": ["A-WALL"],
                "views": ["Level 1"],
                "entities": ["Door 7"],
                "stdout": "sensitive stdout should be redacted",
                "stderr": "sensitive stderr should be redacted",
            },
            command_exit_code=0,
            timeout_seconds=30,
            source_extension=".dwg",
        )

    def get_diagnostics(self) -> dict[str, object]:
        return {"calls": len(self.calls)}


def _sha256_hex(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _write_structural_docx(path: Path) -> None:
    document = docx.Document()
    document.core_properties.title = "Structural Notes"
    document.core_properties.author = "Test Engineer"
    document.add_heading("Structural Notes", level=1)
    document.add_paragraph("S07_DOCX_BEAM_NOTE keeps the docx path searchable.")
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def _write_loads_workbook(path: Path) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Loads"
    sheet["A1"] = "Member"
    sheet["B1"] = "Axial [kN]"
    sheet["A2"] = "Beam-7"
    sheet["B2"] = "S07_XLSX_AXIAL_UNIQUE_42"
    path.parent.mkdir(parents=True, exist_ok=True)
    workbook.save(path)


def _record_with_content(records: list[dict[str, object]], needle: str) -> dict[str, object]:
    return next(record for record in records if needle in str(record["content"]))


async def _assert_recall_hit(tool, query: str, record: dict[str, object]) -> None:
    result = await tool.ainvoke({"query": query, "k": 5})
    assert result == f"- ({record['id']}) {record['content']}"


class TestEngineeringProjectIngestion:
    async def test_folder_ingestion_preserves_partial_failures_and_inline_recall_provenance(
        self,
        tmp_path: Path,
        monkeypatch: pytest.MonkeyPatch,
    ) -> None:
        engineering_dir = tmp_path / "engineering"
        specs_dir = engineering_dir / "specs"
        calcs_dir = engineering_dir / "calcs"
        drawings_dir = engineering_dir / "drawings"
        archive_dir = engineering_dir / "archive"
        hidden_dir = tmp_path / ".git"
        converted_dir = tmp_path / "converted"

        structural_docx = specs_dir / "structural.docx"
        loads_xlsx = calcs_dir / "loads.xlsx"
        bad_xlsx = calcs_dir / "bad.xlsx"
        site_dwg = drawings_dir / "site.dwg"
        archived_text = archive_dir / "old.txt"
        mystery_bin = tmp_path / "mystery.bin"
        hidden_file = hidden_dir / "ignored.txt"
        converted_pdf = converted_dir / "site.pdf"

        _write_structural_docx(structural_docx)
        _write_loads_workbook(loads_xlsx)
        bad_xlsx.parent.mkdir(parents=True, exist_ok=True)
        bad_xlsx.write_bytes(b"not a real workbook")
        site_dwg.parent.mkdir(parents=True, exist_ok=True)
        site_dwg.write_bytes(b"dwg body")
        archived_text.parent.mkdir(parents=True, exist_ok=True)
        archived_text.write_text("archived text", encoding="utf-8")
        mystery_bin.write_bytes(b"\x00\x01\x02")
        hidden_file.parent.mkdir(parents=True, exist_ok=True)
        hidden_file.write_text("hidden text", encoding="utf-8")

        converter = RecordingEngineeringConverter(site_dwg, converted_pdf)

        class FakePdfReader:
            def __init__(self, reader_path: str) -> None:
                assert reader_path == str(converted_pdf)
                self.pages = [
                    FakePage(
                        "\n".join(
                            [
                                "Label: S07_DWG_GRID_A7",
                                "Layer: A-WALL",
                                "View: Level 1",
                                "Entity: Door 7",
                            ]
                        )
                    )
                ]

        monkeypatch.setattr(converted_drawing_elements, "PdfReader", FakePdfReader)

        fake_kb = FakeKB()
        async with lifespan_document_registry(":memory:") as registry:
            result = await ingest_directory(
                fake_kb,
                registry,
                str(tmp_path),
                engineering_converter=converter,
            )

            assert result.ingested_files == 3
            assert result.ingested_chunks == 14
            assert len(result.memory_ids) == 14
            assert len(set(result.memory_ids)) == 14
            assert converter.calls == [str(site_dwg)]

            docx_row = registry.get_by_hash(_sha256_hex(structural_docx))
            assert docx_row is not None
            assert docx_row.status == "indexed"
            assert docx_row.error is None
            assert docx_row.memory_ids
            assert Path(docx_row.stored_path) == structural_docx

            xlsx_row = registry.get_by_hash(_sha256_hex(loads_xlsx))
            assert xlsx_row is not None
            assert xlsx_row.status == "indexed"
            assert xlsx_row.error is None
            assert xlsx_row.memory_ids
            assert Path(xlsx_row.stored_path) == loads_xlsx

            drawing_row = registry.get_by_hash(_sha256_hex(site_dwg))
            assert drawing_row is not None
            assert drawing_row.status == "indexed"
            assert drawing_row.error is None
            assert drawing_row.memory_ids
            assert Path(drawing_row.stored_path) == site_dwg

            bad_row = registry.get_by_hash(_sha256_hex(bad_xlsx))
            assert bad_row is not None
            assert bad_row.status == "failed"
            assert bad_row.error
            assert bad_row.memory_ids == []

            archive_row = registry.get_by_hash(_sha256_hex(archived_text))
            assert archive_row is not None
            assert archive_row.status == "skipped"
            assert archive_row.error == "backup_or_temp"
            assert archive_row.memory_ids == []

            mystery_row = registry.get_by_hash(_sha256_hex(mystery_bin))
            assert mystery_row is not None
            assert mystery_row.status == "skipped"
            assert mystery_row.error == "unsupported_extension"
            assert mystery_row.memory_ids == []

            hidden_row = registry.get_by_hash(_sha256_hex(hidden_file))
            assert hidden_row is None

        memories = fake_kb.dump()
        assert len(memories) == 14
        assert {record["id"] for record in memories} == set(result.memory_ids)

        docx_memories = [
            record for record in memories if record["metadata"]["source"] == "structural.docx"
        ]
        xlsx_memories = [
            record for record in memories if record["metadata"]["source"] == "loads.xlsx"
        ]
        drawing_memories = [
            record for record in memories if record["metadata"]["source"] == "site.dwg"
        ]

        assert len(docx_memories) == 3
        assert len(xlsx_memories) == 6
        assert len(drawing_memories) == 5
        assert set(docx_row.memory_ids) == {record["id"] for record in docx_memories}
        assert set(xlsx_row.memory_ids) == {record["id"] for record in xlsx_memories}
        assert set(drawing_row.memory_ids) == {record["id"] for record in drawing_memories}

        docx_paragraph = _record_with_content(memories, "S07_DOCX_BEAM_NOTE")
        assert docx_paragraph["content"] == (
            "[source=structural.docx; section=Structural Notes; element=paragraph; "
            "extraction=docx_paragraph]\n"
            "S07_DOCX_BEAM_NOTE keeps the docx path searchable."
        )
        assert docx_paragraph["metadata"] == {
            "document_id": docx_row.document_id,
            "source": "structural.docx",
            "path": str(structural_docx),
            "page": None,
            "element_type": "paragraph",
            "extraction_mode": "docx_paragraph",
            "warnings": [],
            "subject": "engineering_narrative",
            "block_index": 1,
            "style_name": "Normal",
            "section_heading": "Structural Notes",
        }

        xlsx_sheet_summary = next(
            record
            for record in xlsx_memories
            if record["metadata"]["extraction_mode"] == "xlsx_sheet_summary"
        )
        assert xlsx_sheet_summary["content"].startswith(
            "[source=loads.xlsx; sheet=Loads; range=A1:B2; element=sheet_summary; "
            "extraction=xlsx_sheet_summary; confidence=1.0]\n"
        )
        assert "Range: A1:B2" in str(xlsx_sheet_summary["content"])
        assert xlsx_sheet_summary["metadata"] == {
            "document_id": xlsx_row.document_id,
            "source": "loads.xlsx",
            "path": str(loads_xlsx),
            "page": None,
            "element_type": "sheet_summary",
            "extraction_mode": "xlsx_sheet_summary",
            "confidence": 1.0,
            "warnings": [],
            "subject": "engineering_workbook",
            "xlsx_sheet": "Loads",
            "xlsx_range": "A1:B2",
            "xlsx_sheet_state": "visible",
            "xlsx_non_empty_cell_count": 4,
            "xlsx_formula_cell_count": 0,
            "xlsx_comment_count": 0,
        }

        xlsx_cell = _record_with_content(memories, "S07_XLSX_AXIAL_UNIQUE_42")
        assert xlsx_cell["content"] == (
            "[source=loads.xlsx; sheet=Loads; cell=B2; label=Beam-7; unit=kN; element=cell; "
            "extraction=xlsx_cell; certainty=exact; confidence=1.0]\n"
            "Sheet: Loads\n"
            "Cell: B2\n"
            "Value: S07_XLSX_AXIAL_UNIQUE_42"
        )
        assert xlsx_cell["metadata"] == {
            "document_id": xlsx_row.document_id,
            "source": "loads.xlsx",
            "path": str(loads_xlsx),
            "page": None,
            "element_type": "cell",
            "extraction_mode": "xlsx_cell",
            "confidence": 1.0,
            "warnings": [],
            "subject": "engineering_workbook",
            "xlsx_sheet": "Loads",
            "xlsx_cell": "B2",
            "xlsx_sheet_state": "visible",
            "xlsx_row_label": "Beam-7",
            "xlsx_column_label": "Axial [kN]",
            "xlsx_label": "Beam-7",
            "xlsx_unit": "kN",
            "xlsx_value": "S07_XLSX_AXIAL_UNIQUE_42",
            "xlsx_value_kind": "literal",
            "extraction_certainty": "exact",
        }

        drawing_summary = next(
            record
            for record in drawing_memories
            if record["metadata"]["extraction_mode"] == "converted_drawing_text_summary"
        )
        assert drawing_summary["content"].startswith(
            "[source=site.dwg; layers=A-WALL; views=Level 1; entities=Door 7; "
            "element=drawing; extraction=converted_drawing_text_summary; confidence=1.0; "
            "warnings=converter_note]\n"
        )
        assert "Source CAD path: " + str(site_dwg) in str(drawing_summary["content"])
        assert "Derived artifact path: " + str(converted_pdf) in str(drawing_summary["content"])
        assert "Conversion warnings: converter_note" in str(drawing_summary["content"])
        assert "stdout" not in drawing_summary["metadata"]["conversion_diagnostics"]
        assert "stderr" not in drawing_summary["metadata"]["conversion_diagnostics"]
        assert drawing_summary["metadata"] == {
            "document_id": drawing_row.document_id,
            "source": "site.dwg",
            "path": str(converted_pdf),
            "page": None,
            "element_type": "drawing",
            "extraction_mode": "converted_drawing_text_summary",
            "confidence": 1.0,
            "warnings": ["converter_note"],
            "subject": "converted_drawing",
            "source_cad_file": "site.dwg",
            "source_cad_path": str(site_dwg),
            "derived_artifact_path": str(converted_pdf),
            "conversion_status": "success",
            "conversion_source_extension": ".dwg",
            "conversion_warnings": ["converter_note"],
            "drawing_artifact_extension": ".pdf",
            "conversion_diagnostics": {
                "layers": ["A-WALL"],
                "views": ["Level 1"],
                "entities": ["Door 7"],
            },
            "drawing_fact_type": "summary",
            "drawing_page_count": 1,
            "drawing_text_page_count": 1,
            "drawing_fact_count": 4,
            "drawing_fact_types": ["label", "layer", "entity_view"],
            "drawing_layers": ["A-WALL"],
            "drawing_views": ["Level 1"],
            "drawing_entities": ["Door 7"],
        }

        drawing_fact = _record_with_content(memories, "S07_DWG_GRID_A7")
        assert drawing_fact["content"] == (
            "[source=site.dwg; page=1; layers=A-WALL; views=Level 1; entities=Door 7; "
            "fact=label; label=S07_DWG_GRID_A7; line=1; element=drawing_fact; "
            "extraction=converted_drawing_text_fact; confidence=1.0; warnings=converter_note]\n"
            "Label: S07_DWG_GRID_A7"
        )
        assert drawing_fact["metadata"] == {
            "document_id": drawing_row.document_id,
            "source": "site.dwg",
            "path": str(converted_pdf),
            "page": 1,
            "element_type": "drawing_fact",
            "extraction_mode": "converted_drawing_text_fact",
            "confidence": 1.0,
            "warnings": ["converter_note"],
            "subject": "converted_drawing",
            "source_cad_file": "site.dwg",
            "source_cad_path": str(site_dwg),
            "derived_artifact_path": str(converted_pdf),
            "conversion_status": "success",
            "conversion_source_extension": ".dwg",
            "conversion_warnings": ["converter_note"],
            "drawing_artifact_extension": ".pdf",
            "conversion_diagnostics": {
                "layers": ["A-WALL"],
                "views": ["Level 1"],
                "entities": ["Door 7"],
            },
            "drawing_fact_type": "label",
            "drawing_fact_value": "S07_DWG_GRID_A7",
            "drawing_line_number": 1,
            "drawing_layers": ["A-WALL"],
            "drawing_views": ["Level 1"],
            "drawing_entities": ["Door 7"],
        }

        kb_recall = next(tool for tool in build_kb_tools(fake_kb) if tool.name == "kb_recall")
        await _assert_recall_hit(kb_recall, "S07_DOCX_BEAM_NOTE", docx_paragraph)
        await _assert_recall_hit(kb_recall, "S07_XLSX_AXIAL_UNIQUE_42", xlsx_cell)
        await _assert_recall_hit(kb_recall, "S07_DWG_GRID_A7", drawing_fact)

    async def test_uploads_via_live_route_preserve_partial_failures_and_inline_recall_provenance(
        self,
        client: TestClient,
        fake_kb: FakeKB,
        tmp_path: Path,
        monkeypatch: pytest.MonkeyPatch,
    ) -> None:
        source_dir = tmp_path / "source"
        documents_dir = tmp_path / "documents"
        converted_dir = tmp_path / "converted"
        structural_docx = source_dir / "structural.docx"
        loads_xlsx = source_dir / "loads.xlsx"
        _write_structural_docx(structural_docx)
        _write_loads_workbook(loads_xlsx)

        structural_docx_body = structural_docx.read_bytes()
        loads_xlsx_body = loads_xlsx.read_bytes()
        dwg_body = b"dwg body"
        bad_xlsx_body = b"not a real workbook"
        temp_docx_body = b"temp draft content"
        mystery_body = b"\x00\x01\x02"
        converted_path = converted_dir / "site.pdf"

        class LiveRouteEngineeringConverter:
            def __init__(self, converted_path: Path) -> None:
                self._converted_path = converted_path
                self.calls: list[str] = []

            def convert(self, source_path: str) -> ConversionResult:
                self.calls.append(source_path)
                self._converted_path.parent.mkdir(parents=True, exist_ok=True)
                self._converted_path.write_bytes(b"%PDF-1.7\n")
                return ConversionResult(
                    success=True,
                    status="success",
                    output_path=str(self._converted_path),
                    warnings=("converter_note",),
                    error=None,
                    diagnostics={
                        "layers": ["A-WALL"],
                        "views": ["Level 1"],
                        "entities": ["Door 7"],
                        "stdout": "sensitive stdout should be redacted",
                        "stderr": "sensitive stderr should be redacted",
                    },
                    command_exit_code=0,
                    timeout_seconds=30,
                    source_extension=".dwg",
                )

            def get_diagnostics(self) -> dict[str, object]:
                return {"calls": len(self.calls)}

        converter = LiveRouteEngineeringConverter(converted_path)

        class FakePdfReader:
            def __init__(self, reader_path: str) -> None:
                assert reader_path == str(converted_path)
                self.pages = [
                    FakePage(
                        "\n".join(
                            [
                                "Label: S07_DWG_GRID_A7",
                                "Layer: A-WALL",
                                "View: Level 1",
                                "Entity: Door 7",
                            ]
                        )
                    )
                ]

        monkeypatch.setattr(converted_drawing_elements, "PdfReader", FakePdfReader)
        client.app.state.app_state.settings.documents_dir = str(documents_dir)
        client.app.state.app_state.engineering_converter = converter
        client.app.state.app_state.engineering_converter_output_dir = str(converted_dir)
        kb = client.app.state.app_state.kb
        assert kb is fake_kb

        response = client.post(
            "/api/ingest",
            files=[
                (
                    "files",
                    (
                        "structural.docx",
                        io.BytesIO(structural_docx_body),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                ),
                (
                    "files",
                    (
                        "loads.xlsx",
                        io.BytesIO(loads_xlsx_body),
                        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    ),
                ),
                ("files", ("site.dwg", io.BytesIO(dwg_body), "application/acad")),
                (
                    "files",
                    (
                        "bad.xlsx",
                        io.BytesIO(bad_xlsx_body),
                        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    ),
                ),
                (
                    "files",
                    (
                        "~$draft.docx",
                        io.BytesIO(temp_docx_body),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                ),
                (
                    "files",
                    ("mystery.bin", io.BytesIO(mystery_body), "application/octet-stream"),
                ),
            ],
        )

        assert response.status_code == 200, response.text
        body = response.json()
        assert body["ingested_files"] == 3
        assert body["ingested_chunks"] == 14
        assert len(body["memory_ids"]) == 14
        assert len(set(body["memory_ids"])) == 14

        docx_record = client.app.state.app_state.registry.get_by_hash(
            hashlib.sha256(structural_docx_body).hexdigest()
        )
        xlsx_record = client.app.state.app_state.registry.get_by_hash(
            hashlib.sha256(loads_xlsx_body).hexdigest()
        )
        drawing_record = client.app.state.app_state.registry.get_by_hash(
            hashlib.sha256(dwg_body).hexdigest()
        )
        bad_xlsx_record = client.app.state.app_state.registry.get_by_hash(
            hashlib.sha256(bad_xlsx_body).hexdigest()
        )
        temp_record = client.app.state.app_state.registry.get_by_hash(
            hashlib.sha256(temp_docx_body).hexdigest()
        )
        mystery_record = client.app.state.app_state.registry.get_by_hash(
            hashlib.sha256(mystery_body).hexdigest()
        )

        assert docx_record is not None
        assert xlsx_record is not None
        assert drawing_record is not None
        assert bad_xlsx_record is not None
        assert temp_record is not None
        assert mystery_record is not None

        assert docx_record.original_filename == "structural.docx"
        assert docx_record.status == "indexed"
        assert docx_record.error is None
        assert len(docx_record.memory_ids) == 3
        assert Path(docx_record.stored_path).parent == documents_dir
        assert Path(docx_record.stored_path).name == f"{docx_record.document_id}.docx"
        assert Path(docx_record.stored_path).read_bytes() == structural_docx_body

        assert xlsx_record.original_filename == "loads.xlsx"
        assert xlsx_record.status == "indexed"
        assert xlsx_record.error is None
        assert len(xlsx_record.memory_ids) == 6
        assert Path(xlsx_record.stored_path).parent == documents_dir
        assert Path(xlsx_record.stored_path).name == f"{xlsx_record.document_id}.xlsx"
        assert Path(xlsx_record.stored_path).read_bytes() == loads_xlsx_body

        assert drawing_record.original_filename == "site.dwg"
        assert drawing_record.status == "indexed"
        assert drawing_record.error is None
        assert len(drawing_record.memory_ids) == 5
        assert Path(drawing_record.stored_path).parent == documents_dir
        assert Path(drawing_record.stored_path).name == f"{drawing_record.document_id}.dwg"
        assert Path(drawing_record.stored_path).read_bytes() == dwg_body
        assert converter.calls == [drawing_record.stored_path]

        assert bad_xlsx_record.original_filename == "bad.xlsx"
        assert bad_xlsx_record.status == "failed"
        assert bad_xlsx_record.error
        assert bad_xlsx_record.memory_ids == []
        assert Path(bad_xlsx_record.stored_path).parent == documents_dir
        assert Path(bad_xlsx_record.stored_path).name == f"{bad_xlsx_record.document_id}.xlsx"
        assert Path(bad_xlsx_record.stored_path).read_bytes() == bad_xlsx_body

        assert temp_record.original_filename == "~$draft.docx"
        assert temp_record.status == "skipped"
        assert temp_record.error == "backup_or_temp"
        assert temp_record.memory_ids == []
        assert Path(temp_record.stored_path).parent == documents_dir
        assert Path(temp_record.stored_path).name == f"{temp_record.document_id}.docx"
        assert Path(temp_record.stored_path).read_bytes() == temp_docx_body

        assert mystery_record.original_filename == "mystery.bin"
        assert mystery_record.status == "skipped"
        assert mystery_record.error == "unsupported_extension"
        assert mystery_record.memory_ids == []
        assert Path(mystery_record.stored_path).parent == documents_dir
        assert Path(mystery_record.stored_path).name == f"{mystery_record.document_id}.bin"
        assert Path(mystery_record.stored_path).read_bytes() == mystery_body

        assert body["memory_ids"] == [
            *docx_record.memory_ids,
            *xlsx_record.memory_ids,
            *drawing_record.memory_ids,
        ]
        assert [record["id"] for record in fake_kb.dump()] == body["memory_ids"]

        memories = fake_kb.dump()
        assert len(memories) == 14
        assert {record["metadata"]["source"] for record in memories} == {
            "structural.docx",
            "loads.xlsx",
            "site.dwg",
        }

        docx_memories = [
            record for record in memories if record["metadata"]["source"] == "structural.docx"
        ]
        xlsx_memories = [
            record for record in memories if record["metadata"]["source"] == "loads.xlsx"
        ]
        drawing_memories = [
            record for record in memories if record["metadata"]["source"] == "site.dwg"
        ]
        assert len(docx_memories) == 3
        assert len(xlsx_memories) == 6
        assert len(drawing_memories) == 5

        docx_paragraph = _record_with_content(memories, "S07_DOCX_BEAM_NOTE")
        assert docx_paragraph["content"] == (
            "[source=structural.docx; section=Structural Notes; element=paragraph; "
            "extraction=docx_paragraph]\n"
            "S07_DOCX_BEAM_NOTE keeps the docx path searchable."
        )
        assert docx_paragraph["metadata"] == {
            "document_id": docx_record.document_id,
            "source": "structural.docx",
            "path": docx_record.stored_path,
            "page": None,
            "element_type": "paragraph",
            "extraction_mode": "docx_paragraph",
            "warnings": [],
            "subject": "engineering_narrative",
            "block_index": 1,
            "style_name": "Normal",
            "section_heading": "Structural Notes",
        }

        xlsx_cell = _record_with_content(memories, "S07_XLSX_AXIAL_UNIQUE_42")
        assert xlsx_cell["content"] == (
            "[source=loads.xlsx; sheet=Loads; cell=B2; label=Beam-7; unit=kN; element=cell; "
            "extraction=xlsx_cell; certainty=exact; confidence=1.0]\n"
            "Sheet: Loads\n"
            "Cell: B2\n"
            "Value: S07_XLSX_AXIAL_UNIQUE_42"
        )
        assert xlsx_cell["metadata"] == {
            "document_id": xlsx_record.document_id,
            "source": "loads.xlsx",
            "path": xlsx_record.stored_path,
            "page": None,
            "element_type": "cell",
            "extraction_mode": "xlsx_cell",
            "confidence": 1.0,
            "warnings": [],
            "subject": "engineering_workbook",
            "xlsx_sheet": "Loads",
            "xlsx_cell": "B2",
            "xlsx_sheet_state": "visible",
            "xlsx_row_label": "Beam-7",
            "xlsx_column_label": "Axial [kN]",
            "xlsx_label": "Beam-7",
            "xlsx_unit": "kN",
            "xlsx_value": "S07_XLSX_AXIAL_UNIQUE_42",
            "xlsx_value_kind": "literal",
            "extraction_certainty": "exact",
        }

        xlsx_sheet_summary = next(
            record
            for record in xlsx_memories
            if record["metadata"]["extraction_mode"] == "xlsx_sheet_summary"
        )
        assert xlsx_sheet_summary["content"].startswith(
            "[source=loads.xlsx; sheet=Loads; range=A1:B2; element=sheet_summary; "
            "extraction=xlsx_sheet_summary; confidence=1.0]\n"
        )
        assert "Range: A1:B2" in str(xlsx_sheet_summary["content"])

        drawing_summary = next(
            record
            for record in drawing_memories
            if record["metadata"]["extraction_mode"] == "converted_drawing_text_summary"
        )
        assert drawing_summary["content"].startswith(
            "[source=site.dwg; layers=A-WALL; views=Level 1; entities=Door 7; "
            "element=drawing; extraction=converted_drawing_text_summary; confidence=1.0; "
            "warnings=converter_note]\n"
        )
        assert "Source CAD path: " + drawing_record.stored_path in str(drawing_summary["content"])
        assert "Derived artifact path: " + str(converted_path) in str(drawing_summary["content"])
        assert "Conversion warnings: converter_note" in str(drawing_summary["content"])
        assert drawing_summary["metadata"]["conversion_diagnostics"] == {
            "layers": ["A-WALL"],
            "views": ["Level 1"],
            "entities": ["Door 7"],
        }
        assert "stdout" not in drawing_summary["metadata"]["conversion_diagnostics"]
        assert "stderr" not in drawing_summary["metadata"]["conversion_diagnostics"]

        drawing_fact = _record_with_content(memories, "S07_DWG_GRID_A7")
        assert drawing_fact["content"] == (
            "[source=site.dwg; page=1; layers=A-WALL; views=Level 1; entities=Door 7; "
            "fact=label; label=S07_DWG_GRID_A7; line=1; element=drawing_fact; "
            "extraction=converted_drawing_text_fact; confidence=1.0; warnings=converter_note]\n"
            "Label: S07_DWG_GRID_A7"
        )
        assert drawing_fact["metadata"] == {
            "document_id": drawing_record.document_id,
            "source": "site.dwg",
            "path": str(converted_path),
            "page": 1,
            "element_type": "drawing_fact",
            "extraction_mode": "converted_drawing_text_fact",
            "confidence": 1.0,
            "warnings": ["converter_note"],
            "subject": "converted_drawing",
            "source_cad_file": "site.dwg",
            "source_cad_path": drawing_record.stored_path,
            "derived_artifact_path": str(converted_path),
            "conversion_status": "success",
            "conversion_source_extension": ".dwg",
            "conversion_warnings": ["converter_note"],
            "drawing_artifact_extension": ".pdf",
            "conversion_diagnostics": {
                "layers": ["A-WALL"],
                "views": ["Level 1"],
                "entities": ["Door 7"],
            },
            "drawing_fact_type": "label",
            "drawing_fact_value": "S07_DWG_GRID_A7",
            "drawing_line_number": 1,
            "drawing_layers": ["A-WALL"],
            "drawing_views": ["Level 1"],
            "drawing_entities": ["Door 7"],
        }

        kb_recall = next(tool for tool in build_kb_tools(kb) if tool.name == "kb_recall")
        await _assert_recall_hit(kb_recall, "S07_DOCX_BEAM_NOTE", docx_paragraph)
        await _assert_recall_hit(kb_recall, "S07_XLSX_AXIAL_UNIQUE_42", xlsx_cell)
        await _assert_recall_hit(kb_recall, "S07_DWG_GRID_A7", drawing_fact)
