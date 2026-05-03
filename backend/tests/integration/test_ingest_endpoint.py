"""POST /api/ingest accepts multipart uploads and stores them in the KB."""

from __future__ import annotations

import hashlib
import io
from pathlib import Path

import docx
import pytest
from fastapi.testclient import TestClient
from openpyxl import Workbook
from openpyxl.comments import Comment
from openpyxl.worksheet.table import Table

import app.services.converted_drawing_elements as converted_drawing_elements
from app.agent.tools import build_kb_tools
from app.config import Settings
from app.kb.fake import FakeKB
from app.services import ingestion as ingestion_module
from app.services.document_analysis import (
    DocumentAnalysisRefusalError,
    VisualEnrichmentOutput,
    build_document_analyzer,
)
from app.services.document_elements import DocumentElement
from app.services.engineering_converters import ConversionResult
from app.services.ocr_elements import ocr_element_from_text
from app.services.table_elements import RAGGED_TABLE_WARNING, table_element_from_rows
from app.services.visual_elements import APPROXIMATE_VALUE_WARNING, visual_element_from_summary


def _registry_row_count_for_hash(client: TestClient, content_hash: str) -> int:
    row = client.app.state.app_state.registry._conn.execute(
        "SELECT COUNT(*) AS count FROM documents WHERE content_hash = ?",
        (content_hash,),
    ).fetchone()
    assert row is not None
    return int(row["count"])


class RecordingDocumentAnalysisClient:
    def __init__(
        self,
        *,
        responses: dict[str, VisualEnrichmentOutput] | None = None,
        exception: Exception | None = None,
    ) -> None:
        self.responses = responses or {}
        self.exception = exception
        self.calls: list[DocumentElement] = []

    def enrich(self, element: DocumentElement) -> VisualEnrichmentOutput:
        self.calls.append(element)
        if self.exception is not None:
            raise self.exception
        response = self.responses.get(element.element_type)
        assert response is not None
        return response


def _build_visual_analyzer(
    *,
    responses: dict[str, VisualEnrichmentOutput] | None = None,
    exception: Exception | None = None,
):
    analysis_client = RecordingDocumentAnalysisClient(
        responses=responses,
        exception=exception,
    )
    analyzer = build_document_analyzer(
        Settings(
            llm_provider="ollama",
            document_analysis_enabled=True,
            document_analysis_api_key="",
            document_analysis_model="fake-visual-model",
        ),
        client=analysis_client,
    )
    return analyzer, analysis_client


class RecordingEngineeringConverter:
    def __init__(self, convert) -> None:
        self.calls: list[str] = []
        self._convert = convert

    def convert(self, source_path: str) -> ConversionResult:
        self.calls.append(source_path)
        return self._convert(source_path)

    def get_diagnostics(self) -> dict[str, object]:
        return {"calls": len(self.calls)}


def _build_xlsx_workbook_bytes() -> bytes:
    workbook = Workbook()
    visible_sheet = workbook.active
    visible_sheet.title = "Loads"
    visible_sheet["A1"] = "Region"
    visible_sheet["B1"] = "Load [kN]"
    visible_sheet["A2"] = "North [kN]"
    visible_sheet["B2"] = 12
    visible_sheet["C2"] = "=SUM(B2:B2)"
    visible_sheet["C2"].comment = Comment("Needs review", "Planner")
    visible_sheet.add_table(Table(displayName="LoadTable", ref="A1:B2"))

    hidden_sheet = workbook.create_sheet("Hidden Notes")
    hidden_sheet["A1"] = "Internal note"
    hidden_sheet.sheet_state = "hidden"

    buffer = io.BytesIO()
    workbook.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


class TestIngestEndpoint:
    @staticmethod
    def _make_converted_result(source_path: str, converted_path: Path) -> ConversionResult:
        converted_path.parent.mkdir(parents=True, exist_ok=True)
        converted_path.write_text("converted north drawing", encoding="utf-8")
        return ConversionResult(
            success=True,
            status="success",
            output_path=str(converted_path),
            warnings=("converter_note",),
            error=None,
            diagnostics={
                "fake": True,
                "source_path": source_path,
                "layers": ["A-WALL"],
                "views": ["Level 1"],
                "entities": ["Door 7"],
                "stdout": "converter stdout should stay hidden",
            },
            command_exit_code=0,
            timeout_seconds=30,
            source_extension=Path(source_path).suffix.lower(),
        )

    def test_uploads_text_file(self, client: TestClient, tmp_path: Path) -> None:
        # Redirect documents_dir into the test tmpdir so we don't write
        # outside the sandbox.
        client.app.state.app_state.settings.documents_dir = str(tmp_path)

        body_bytes = b"hello world"
        files = [("files", ("note.txt", io.BytesIO(body_bytes), "text/plain"))]
        r = client.post("/api/ingest", files=files)
        assert r.status_code == 200, r.text
        body = r.json()
        assert body["ingested_files"] == 1
        assert body["ingested_chunks"] == 1
        assert len(body["memory_ids"]) == 1

        record = client.app.state.app_state.registry.get_by_hash(
            hashlib.sha256(body_bytes).hexdigest()
        )
        assert record is not None
        assert record.original_filename == "note.txt"
        assert record.status == "indexed"
        assert record.error is None
        assert record.memory_ids == body["memory_ids"]
        assert Path(record.stored_path).parent == tmp_path
        assert Path(record.stored_path).name == f"{record.document_id}.txt"
        assert Path(record.stored_path).exists()
        assert not (tmp_path / "note.txt").exists()

    def test_duplicate_upload_reuses_registry_memory_ids_without_new_kb_records(
        self,
        client: TestClient,
        fake_kb: FakeKB,
        tmp_path: Path,
    ) -> None:
        client.app.state.app_state.settings.documents_dir = str(tmp_path)
        body_bytes = b"same content"
        content_hash = hashlib.sha256(body_bytes).hexdigest()

        first = client.post(
            "/api/ingest",
            files=[("files", ("first.txt", io.BytesIO(body_bytes), "text/plain"))],
        )
        assert first.status_code == 200, first.text
        records_after_first = fake_kb.dump()
        assert len(records_after_first) == 1
        first_record = client.app.state.app_state.registry.get_by_hash(content_hash)
        assert first_record is not None

        second = client.post(
            "/api/ingest",
            files=[("files", ("second.txt", io.BytesIO(body_bytes), "text/plain"))],
        )
        assert second.status_code == 200, second.text
        assert second.json()["ingested_files"] == 0
        assert second.json()["ingested_chunks"] == 0
        assert second.json()["memory_ids"] == first.json()["memory_ids"]
        assert fake_kb.dump() == records_after_first
        assert _registry_row_count_for_hash(client, content_hash) == 1

        record = client.app.state.app_state.registry.get_by_hash(content_hash)
        assert record is not None
        assert record.document_id == first_record.document_id
        assert record.original_filename == "first.txt"
        assert record.memory_ids == first.json()["memory_ids"]
        assert len(list(tmp_path.iterdir())) == 1

    def test_same_filename_different_content_creates_distinct_registry_rows_and_paths(
        self,
        client: TestClient,
        tmp_path: Path,
    ) -> None:
        client.app.state.app_state.settings.documents_dir = str(tmp_path)
        first_bytes = b"first drawing revision"
        second_bytes = b"second drawing revision"

        first = client.post(
            "/api/ingest",
            files=[("files", ("plan.txt", io.BytesIO(first_bytes), "text/plain"))],
        )
        second = client.post(
            "/api/ingest",
            files=[("files", ("plan.txt", io.BytesIO(second_bytes), "text/plain"))],
        )

        assert first.status_code == 200, first.text
        assert second.status_code == 200, second.text
        first_record = client.app.state.app_state.registry.get_by_hash(
            hashlib.sha256(first_bytes).hexdigest()
        )
        second_record = client.app.state.app_state.registry.get_by_hash(
            hashlib.sha256(second_bytes).hexdigest()
        )
        assert first_record is not None
        assert second_record is not None
        assert first_record.document_id != second_record.document_id
        assert first_record.original_filename == "plan.txt"
        assert second_record.original_filename == "plan.txt"
        assert first_record.stored_path != second_record.stored_path
        assert Path(first_record.stored_path).parent == tmp_path
        assert Path(second_record.stored_path).parent == tmp_path
        assert Path(first_record.stored_path).exists()
        assert Path(second_record.stored_path).exists()

    def test_same_batch_duplicate_reuses_newly_persisted_memory_ids(
        self,
        client: TestClient,
        fake_kb: FakeKB,
        tmp_path: Path,
    ) -> None:
        client.app.state.app_state.settings.documents_dir = str(tmp_path)
        body_bytes = b"same batch content"

        response = client.post(
            "/api/ingest",
            files=[
                ("files", ("first.txt", io.BytesIO(body_bytes), "text/plain")),
                ("files", ("second.txt", io.BytesIO(body_bytes), "text/plain")),
            ],
        )

        assert response.status_code == 200, response.text
        body = response.json()
        assert body["ingested_files"] == 1
        assert body["ingested_chunks"] == 1
        assert len(body["memory_ids"]) == 2
        assert body["memory_ids"][0] == body["memory_ids"][1]
        assert len(fake_kb.dump()) == 1
        assert len(list(tmp_path.iterdir())) == 1

    def test_uploaded_duplicate_row_retries_missing_file_write(
        self,
        client: TestClient,
        fake_kb: FakeKB,
        tmp_path: Path,
    ) -> None:
        client.app.state.app_state.settings.documents_dir = str(tmp_path)
        body_bytes = b"retry content"
        content_hash = hashlib.sha256(body_bytes).hexdigest()
        stored_path = tmp_path / "retrydoc.txt"
        record, _ = client.app.state.app_state.registry.register_or_get(
            content_hash,
            original_filename="first.txt",
            stored_path=str(stored_path),
            content_type="text/plain",
            byte_size=len(body_bytes),
            document_id="retrydoc",
        )
        assert record.status == "uploaded"
        assert not stored_path.exists()

        response = client.post(
            "/api/ingest",
            files=[("files", ("second.txt", io.BytesIO(body_bytes), "text/plain"))],
        )

        assert response.status_code == 200, response.text
        assert response.json()["ingested_files"] == 1
        assert len(fake_kb.dump()) == 1
        updated = client.app.state.app_state.registry.get_by_id("retrydoc")
        assert updated is not None
        assert updated.status == "indexed"
        assert updated.memory_ids == response.json()["memory_ids"]
        assert stored_path.exists()

    def test_rejects_empty_upload(self, client: TestClient) -> None:
        r = client.post("/api/ingest", files=[])
        assert r.status_code in (400, 422)

    def test_rejects_empty_file_body(self, client: TestClient, tmp_path: Path) -> None:
        client.app.state.app_state.settings.documents_dir = str(tmp_path)
        r = client.post(
            "/api/ingest",
            files=[("files", ("empty.txt", io.BytesIO(b""), "text/plain"))],
        )
        assert r.status_code == 400

    def test_unknown_extension_persists_as_skipped(
        self,
        client: TestClient,
        tmp_path: Path,
    ) -> None:
        client.app.state.app_state.settings.documents_dir = str(tmp_path)
        body_bytes = b"hello"

        response = client.post(
            "/api/ingest",
            files=[
                (
                    "files",
                    ("malware.exe", io.BytesIO(body_bytes), "application/octet-stream"),
                )
            ],
        )

        assert response.status_code == 200, response.text
        body = response.json()
        assert body == {"ingested_files": 0, "ingested_chunks": 0, "memory_ids": []}

        record = client.app.state.app_state.registry.get_by_hash(
            hashlib.sha256(body_bytes).hexdigest()
        )
        assert record is not None
        assert record.original_filename == "malware.exe"
        assert record.status == "skipped"
        assert record.error == "unsupported_extension"
        assert record.memory_ids == []
        assert Path(record.stored_path).exists()
        assert Path(record.stored_path).read_bytes() == body_bytes

    @pytest.mark.parametrize(
        ("filename", "content_type", "reason", "body_bytes"),
        [
            pytest.param(
                "north.dwg",
                "application/acad",
                "missing_configuration",
                b"dwg body",
                id="dwg",
            ),
            pytest.param(
                "site.png",
                "image/png",
                "image_extractor_pending",
                b"png body",
                id="png",
            ),
        ],
    )
    def test_engineering_uploads_persist_as_skipped_with_reason(
        self,
        client: TestClient,
        fake_kb: FakeKB,
        tmp_path: Path,
        filename: str,
        content_type: str,
        reason: str,
        body_bytes: bytes,
    ) -> None:
        client.app.state.app_state.settings.documents_dir = str(tmp_path)
        before_dump = fake_kb.dump()

        response = client.post(
            "/api/ingest",
            files=[("files", (filename, io.BytesIO(body_bytes), content_type))],
        )

        assert response.status_code == 200, response.text
        body = response.json()
        assert body == {"ingested_files": 0, "ingested_chunks": 0, "memory_ids": []}
        assert fake_kb.dump() == before_dump

        record = client.app.state.app_state.registry.get_by_hash(
            hashlib.sha256(body_bytes).hexdigest()
        )
        assert record is not None
        assert record.original_filename == filename
        assert record.status == "skipped"
        assert record.error == reason
        assert record.memory_ids == []
        assert Path(record.stored_path).exists()
        assert Path(record.stored_path).read_bytes() == body_bytes

    def test_uploads_cad_export_through_converter_seam(
        self,
        client: TestClient,
        fake_kb: FakeKB,
        tmp_path: Path,
        monkeypatch: pytest.MonkeyPatch,
    ) -> None:
        client.app.state.app_state.settings.documents_dir = str(tmp_path)
        converted_dir = tmp_path / "converted"
        converted_path = converted_dir / "north.pdf"
        converter = RecordingEngineeringConverter(
            lambda source_path: self._make_converted_result(source_path, converted_path)
        )
        client.app.state.app_state.engineering_converter = converter
        client.app.state.app_state.engineering_converter_output_dir = str(converted_dir)

        class FakePage:
            def __init__(self, text: str | None) -> None:
                self._text = text

            def extract_text(self) -> str | None:
                return self._text

        class FakePdfReader:
            def __init__(self, reader_path: str) -> None:
                assert reader_path == str(converted_path)
                self.pages = [
                    FakePage("Label: North entry\nDimension: 12'-0\""),
                    FakePage(
                        "Layer: A-WALL\n"
                        "View: Level 1\n"
                        "Entity: Door 7\n"
                        "Revision: R3\n"
                        "Note: verify field"
                    ),
                ]

        monkeypatch.setattr(converted_drawing_elements, "PdfReader", FakePdfReader)

        def fake_parse_document(*_args: object, **_kwargs: object) -> list[DocumentElement]:
            raise AssertionError("parse_document should not run for converted drawings")

        monkeypatch.setattr(ingestion_module.parsers, "parse_document", fake_parse_document)
        body_bytes = b"dwg body"

        response = client.post(
            "/api/ingest",
            files=[("files", ("north.dwg", io.BytesIO(body_bytes), "application/acad"))],
        )

        assert response.status_code == 200, response.text
        body = response.json()
        assert body["ingested_files"] == 1
        assert body["ingested_chunks"] == len(body["memory_ids"]) == 8
        records = fake_kb.dump()
        assert len(records) == 8
        assert body["memory_ids"] == [record["id"] for record in records]

        registry_record = client.app.state.app_state.registry.get_by_hash(
            hashlib.sha256(body_bytes).hexdigest()
        )
        assert registry_record is not None
        assert registry_record.status == "indexed"
        assert registry_record.error is None
        assert registry_record.memory_ids == body["memory_ids"]
        assert converter.calls == [registry_record.stored_path]
        assert Path(registry_record.stored_path).exists()
        assert Path(converted_path).exists()

        summary_record = next(
            record
            for record in records
            if record["metadata"]["extraction_mode"] == "converted_drawing_text_summary"
        )
        assert summary_record["content"].startswith(
            "[source=north.dwg; layers=A-WALL; views=Level 1; entities=Door 7; "
            "element=drawing; extraction=converted_drawing_text_summary; confidence=1.0; "
            "warnings=converter_note]"
        )
        assert "Text layer mode: exact" in summary_record["content"]
        assert "Conversion warnings: converter_note" in summary_record["content"]
        assert summary_record["metadata"]["source_cad_file"] == "north.dwg"
        assert summary_record["metadata"]["source_cad_path"] == registry_record.stored_path
        assert summary_record["metadata"]["derived_artifact_path"] == str(converted_path)
        assert summary_record["metadata"]["conversion_status"] == "success"
        assert summary_record["metadata"]["conversion_warnings"] == ["converter_note"]
        assert summary_record["metadata"]["drawing_fact_count"] == 7
        assert summary_record["metadata"]["drawing_fact_types"] == [
            "label",
            "dimension",
            "layer",
            "entity_view",
            "revision_marker",
            "visible_note",
        ]
        assert summary_record["metadata"]["drawing_layers"] == ["A-WALL"]
        assert summary_record["metadata"]["drawing_views"] == ["Level 1"]
        assert summary_record["metadata"]["drawing_entities"] == ["Door 7"]
        assert summary_record["metadata"]["conversion_diagnostics"] == {
            "fake": True,
            "source_path": registry_record.stored_path,
            "layers": ["A-WALL"],
            "views": ["Level 1"],
            "entities": ["Door 7"],
        }

        fact_records = [
            record
            for record in records
            if record["metadata"]["extraction_mode"] == "converted_drawing_text_fact"
        ]
        assert len(fact_records) == 7
        assert [record["metadata"]["drawing_fact_type"] for record in fact_records] == [
            "label",
            "dimension",
            "layer",
            "entity_view",
            "entity_view",
            "revision_marker",
            "visible_note",
        ]
        assert [record["metadata"]["drawing_fact_subtype"] for record in fact_records[3:5]] == [
            "view",
            "entity",
        ]
        assert all(
            record["metadata"]["source_cad_path"] == registry_record.stored_path
            for record in fact_records
        )
        assert all(
            record["metadata"]["derived_artifact_path"] == str(converted_path)
            for record in fact_records
        )

    def test_uploads_xlsx_emits_summary_sheet_formula_value_comment_and_table_memories(
        self,
        client: TestClient,
        fake_kb: FakeKB,
        tmp_path: Path,
    ) -> None:
        client.app.state.app_state.settings.documents_dir = str(tmp_path)
        body_bytes = _build_xlsx_workbook_bytes()

        response = client.post(
            "/api/ingest",
            files=[
                (
                    "files",
                    (
                        "loads.xlsx",
                        io.BytesIO(body_bytes),
                        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    ),
                )
            ],
        )

        assert response.status_code == 200, response.text
        body = response.json()
        assert body["ingested_files"] == 1
        assert body["ingested_chunks"] == len(body["memory_ids"])
        assert body["memory_ids"]

        records = fake_kb.dump()
        assert body["memory_ids"] == [record["id"] for record in records]
        assert len(records) == len(body["memory_ids"])

        registry_record = client.app.state.app_state.registry.get_by_hash(
            hashlib.sha256(body_bytes).hexdigest()
        )
        assert registry_record is not None
        assert registry_record.status == "indexed"
        assert registry_record.error is None
        assert registry_record.memory_ids == body["memory_ids"]
        assert Path(registry_record.stored_path).exists()
        assert Path(registry_record.stored_path).parent == tmp_path

        assert all(record["metadata"]["source"] == "loads.xlsx" for record in records)
        assert all(
            record["metadata"]["document_id"] == registry_record.document_id for record in records
        )
        assert all(record["metadata"]["path"] == registry_record.stored_path for record in records)
        assert all(record["content"].startswith("[source=loads.xlsx;") for record in records)

        modes = [record["metadata"]["extraction_mode"] for record in records]
        assert {
            "xlsx_summary",
            "xlsx_sheet_summary",
            "xlsx_cell",
            "xlsx_formula",
            "xlsx_comment",
            "xlsx_table",
        } <= set(modes)
        assert modes.count("xlsx_sheet_summary") == 2

        summary_record = next(
            record for record in records if record["metadata"]["extraction_mode"] == "xlsx_summary"
        )
        assert summary_record["metadata"]["sheet_count"] == 2
        assert summary_record["metadata"]["xlsx_sheets"] == ["Loads", "Hidden Notes"]
        assert summary_record["metadata"]["xlsx_visible_sheet_count"] == 1
        assert summary_record["metadata"]["xlsx_hidden_sheet_count"] == 1
        assert summary_record["content"].startswith(
            "[source=loads.xlsx; element=file_summary; extraction=xlsx_summary; confidence=1.0]"
        )

        sheet_summaries = [
            record
            for record in records
            if record["metadata"]["extraction_mode"] == "xlsx_sheet_summary"
        ]
        assert len(sheet_summaries) == 2
        visible_sheet_summary = next(
            record for record in sheet_summaries if record["metadata"]["xlsx_sheet"] == "Loads"
        )
        hidden_sheet_summary = next(
            record
            for record in sheet_summaries
            if record["metadata"]["xlsx_sheet"] == "Hidden Notes"
        )
        assert visible_sheet_summary["metadata"]["xlsx_sheet_state"] == "visible"
        assert visible_sheet_summary["metadata"]["xlsx_range"] == "A1:C2"
        assert hidden_sheet_summary["metadata"]["xlsx_sheet_state"] == "hidden"

        cell_record = next(
            record
            for record in records
            if record["metadata"]["extraction_mode"] == "xlsx_cell"
            and record["metadata"]["xlsx_cell"] == "B2"
        )
        assert cell_record["content"].startswith(
            "[source=loads.xlsx; sheet=Loads; cell=B2; label=North [kN]; unit=kN; "
            "element=cell; extraction=xlsx_cell; certainty=exact; confidence=1.0]"
        )
        assert cell_record["metadata"]["xlsx_sheet"] == "Loads"
        assert cell_record["metadata"]["xlsx_row_label"] == "North [kN]"
        assert cell_record["metadata"]["xlsx_column_label"] == "Load [kN]"
        assert cell_record["metadata"]["xlsx_label"] == "North [kN]"
        assert cell_record["metadata"]["xlsx_unit"] == "kN"
        assert cell_record["metadata"]["xlsx_value"] == 12
        assert cell_record["metadata"]["xlsx_value_kind"] == "literal"

        formula_record = next(
            record for record in records if record["metadata"]["extraction_mode"] == "xlsx_formula"
        )
        assert formula_record["content"].startswith(
            "[source=loads.xlsx; sheet=Loads; cell=C2; label=12; "
            "element=formula; extraction=xlsx_formula; "
            "certainty=exact_formula_cached_value_unknown; confidence=1.0; "
            "warnings=missing_cached_value]"
        )
        assert formula_record["metadata"]["xlsx_formula"] == "=SUM(B2:B2)"
        assert formula_record["metadata"]["xlsx_value_kind"] == "missing_cached_value"
        assert formula_record["metadata"]["warnings"] == ["missing_cached_value"]

        comment_record = next(
            record for record in records if record["metadata"]["extraction_mode"] == "xlsx_comment"
        )
        assert comment_record["content"].startswith(
            "[source=loads.xlsx; sheet=Loads; cell=C2; "
            "element=comment; extraction=xlsx_comment; certainty=exact; confidence=1.0]"
        )
        assert comment_record["metadata"]["xlsx_comment_author"] == "Planner"
        assert comment_record["metadata"]["xlsx_comment_text"] == "Needs review"

        table_record = next(
            record for record in records if record["metadata"]["extraction_mode"] == "xlsx_table"
        )
        assert table_record["content"].startswith(
            "[source=loads.xlsx; sheet=Loads; range=A1:B2; table=LoadTable; "
            "element=table; extraction=xlsx_table; confidence=1.0]"
        )
        assert table_record["metadata"]["xlsx_sheet"] == "Loads"
        assert table_record["metadata"]["xlsx_range"] == "A1:B2"
        assert table_record["metadata"]["xlsx_table_name"] == "LoadTable"
        assert table_record["metadata"]["table_rows"] == 2
        assert table_record["metadata"]["table_columns"] == 2

    def test_pdf_md_txt_uploads_unaffected_by_classifier(
        self,
        client: TestClient,
        fake_kb: FakeKB,
        tmp_path: Path,
        monkeypatch: pytest.MonkeyPatch,
    ) -> None:
        client.app.state.app_state.settings.documents_dir = str(tmp_path)

        def fake_parse_document(
            path: str,
            *,
            source: str,
            document_id: str | None = None,
        ) -> list[DocumentElement]:
            assert Path(path).exists()
            assert document_id is not None
            return [
                DocumentElement(
                    document_id=document_id,
                    source=source,
                    path=path,
                    page=1,
                    element_type="paragraph",
                    extraction_mode="text",
                    content=f"parsed {source}",
                )
            ]

        monkeypatch.setattr(ingestion_module.parsers, "parse_document", fake_parse_document)

        payloads = [
            ("files", ("design.pdf", io.BytesIO(b"%PDF-1.7\nclassifier"), "application/pdf")),
            ("files", ("notes.md", io.BytesIO(b"# notes"), "text/markdown")),
            ("files", ("readme.txt", io.BytesIO(b"plain text"), "text/plain")),
        ]
        response = client.post("/api/ingest", files=payloads)

        assert response.status_code == 200, response.text
        body = response.json()
        assert body["ingested_files"] == 3
        assert body["ingested_chunks"] == 3
        assert len(body["memory_ids"]) == 3
        assert len(fake_kb.dump()) == 3

        for filename, body_bytes in (
            ("design.pdf", b"%PDF-1.7\nclassifier"),
            ("notes.md", b"# notes"),
            ("readme.txt", b"plain text"),
        ):
            record = client.app.state.app_state.registry.get_by_hash(
                hashlib.sha256(body_bytes).hexdigest()
            )
            assert record is not None
            assert record.original_filename == filename
            assert record.status == "indexed"
            assert record.error is None
            assert record.memory_ids
            assert Path(record.stored_path).exists()

    def test_accepts_uploads_at_and_under_size_limit(
        self,
        client: TestClient,
        tmp_path: Path,
    ) -> None:
        client.app.state.app_state.settings.documents_dir = str(tmp_path)
        client.app.state.app_state.settings.max_upload_bytes = 4

        under_limit = client.post(
            "/api/ingest",
            files=[("files", ("under.txt", io.BytesIO(b"123"), "text/plain"))],
        )
        at_limit = client.post(
            "/api/ingest",
            files=[("files", ("exact.txt", io.BytesIO(b"1234"), "text/plain"))],
        )

        assert under_limit.status_code == 200, under_limit.text
        assert under_limit.json()["ingested_files"] == 1
        assert at_limit.status_code == 200, at_limit.text
        assert at_limit.json()["ingested_files"] == 1

    def test_rejects_oversized_upload(self, client: TestClient, tmp_path: Path) -> None:
        client.app.state.app_state.settings.documents_dir = str(tmp_path)
        client.app.state.app_state.settings.max_upload_bytes = 4
        r = client.post(
            "/api/ingest",
            files=[("files", ("big.txt", io.BytesIO(b"12345"), "text/plain"))],
        )
        assert r.status_code == 413

    def test_parser_failure_marks_registry_failed_and_returns_ingest_response(
        self,
        client: TestClient,
        fake_kb: FakeKB,
        tmp_path: Path,
        monkeypatch: pytest.MonkeyPatch,
    ) -> None:
        client.app.state.app_state.settings.documents_dir = str(tmp_path)

        def raise_boom(*_args: object, **_kwargs: object) -> list[object]:
            raise RuntimeError("boom")

        monkeypatch.setattr(ingestion_module.parsers, "parse_document", raise_boom)
        body_bytes = b"%PDF-1.7\nnot a real pdf"

        response = client.post(
            "/api/ingest",
            files=[("files", ("bad.pdf", io.BytesIO(body_bytes), "application/pdf"))],
        )

        assert response.status_code == 200, response.text
        body = response.json()
        assert body == {"ingested_files": 0, "ingested_chunks": 0, "memory_ids": []}
        record = client.app.state.app_state.registry.get_by_hash(
            hashlib.sha256(body_bytes).hexdigest()
        )
        assert record is not None
        assert record.status == "failed"
        assert record.error == "boom"
        assert record.memory_ids == []
        assert fake_kb.dump() == []

    def test_corrupt_xlsx_upload_marks_registry_failed_and_continues_batch(
        self,
        client: TestClient,
        fake_kb: FakeKB,
        tmp_path: Path,
    ) -> None:
        client.app.state.app_state.settings.documents_dir = str(tmp_path)
        bad_xlsx_body = b"not a real workbook"
        valid_body = b"batch text"

        response = client.post(
            "/api/ingest",
            files=[
                (
                    "files",
                    (
                        "bad.xlsx",
                        io.BytesIO(bad_xlsx_body),
                        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    ),
                ),
                ("files", ("valid.txt", io.BytesIO(valid_body), "text/plain")),
            ],
        )

        assert response.status_code == 200, response.text
        body = response.json()
        assert body["ingested_files"] == 1
        assert body["ingested_chunks"] == 1
        assert len(body["memory_ids"]) == 1

        records = fake_kb.dump()
        assert len(records) == 1
        assert body["memory_ids"] == [record["id"] for record in records]

        text_record = client.app.state.app_state.registry.get_by_hash(
            hashlib.sha256(valid_body).hexdigest()
        )
        assert text_record is not None
        assert text_record.status == "indexed"
        assert text_record.error is None
        assert text_record.memory_ids == body["memory_ids"]
        assert Path(text_record.stored_path).exists()

        xlsx_record = client.app.state.app_state.registry.get_by_hash(
            hashlib.sha256(bad_xlsx_body).hexdigest()
        )
        assert xlsx_record is not None
        assert xlsx_record.status == "failed"
        assert xlsx_record.error
        assert xlsx_record.memory_ids == []
        assert Path(xlsx_record.stored_path).exists()

    def test_uploads_multi_page_pdf_emits_per_page_memories_with_inline_provenance(
        self,
        client: TestClient,
        fake_kb: FakeKB,
        tmp_path: Path,
        monkeypatch: pytest.MonkeyPatch,
    ) -> None:
        client.app.state.app_state.settings.documents_dir = str(tmp_path)

        def fake_parse_document(
            path: str,
            *,
            source: str,
            document_id: str | None = None,
        ) -> list[DocumentElement]:
            assert Path(path).exists()
            assert source == "multi.pdf"
            assert document_id is not None
            return [
                DocumentElement(
                    document_id=document_id,
                    source=source,
                    path=path,
                    page=1,
                    element_type="paragraph",
                    extraction_mode="pdf_text",
                    content="first page body",
                ),
                DocumentElement(
                    document_id=document_id,
                    source=source,
                    path=path,
                    page=2,
                    element_type="paragraph",
                    extraction_mode="pdf_text",
                    content="second page body",
                ),
            ]

        monkeypatch.setattr(ingestion_module.parsers, "parse_document", fake_parse_document)
        body_bytes = b"%PDF-1.7\n"

        response = client.post(
            "/api/ingest",
            files=[("files", ("multi.pdf", io.BytesIO(body_bytes), "application/pdf"))],
        )

        assert response.status_code == 200, response.text
        body = response.json()
        assert body["ingested_files"] == 1
        assert body["ingested_chunks"] == 2
        records = fake_kb.dump()
        assert len(records) == 2

        registry_record = client.app.state.app_state.registry.get_by_hash(
            hashlib.sha256(body_bytes).hexdigest()
        )
        assert registry_record is not None
        assert registry_record.status == "indexed"
        assert registry_record.memory_ids == body["memory_ids"]
        assert body["memory_ids"] == [record["id"] for record in records]

        assert records[0]["content"].startswith("[source=multi.pdf; page=1;")
        assert records[0]["content"].endswith("first page body")
        assert records[1]["content"].startswith("[source=multi.pdf; page=2;")
        assert records[1]["content"].endswith("second page body")
        for record in records:
            metadata = record["metadata"]
            assert metadata["extraction_mode"] == "pdf_text"
            assert metadata["element_type"] == "paragraph"
            assert metadata["document_id"] == registry_record.document_id
        assert [record["metadata"]["page"] for record in records] == [1, 2]

    def test_uploads_visual_summary_elements_as_standalone_memories(
        self,
        client: TestClient,
        fake_kb: FakeKB,
        tmp_path: Path,
        monkeypatch: pytest.MonkeyPatch,
    ) -> None:
        client.app.state.app_state.settings.documents_dir = str(tmp_path)

        def fake_parse_document(
            path: str,
            *,
            source: str,
            document_id: str | None = None,
        ) -> list[DocumentElement]:
            assert Path(path).exists()
            assert source == "visuals.pdf"
            assert document_id is not None
            drawing = visual_element_from_summary(
                "North stair sketch",
                element_type="drawing",
                source=source,
                document_id=document_id,
                path=path,
                page=1,
                confidence=0.88,
                labels=("North stair", "Access path"),
                relationships=("North stair -> Access path",),
                metadata={"figure_index": 0},
            )
            diagram = visual_element_from_summary(
                "Workflow diagram",
                element_type="diagram",
                source=source,
                document_id=document_id,
                path=path,
                page=2,
                confidence=0.74,
                labels=("Parser", "Memory KB"),
                relationships=("Parser -> Ingest", "Ingest -> KB"),
                metadata={"figure_index": 1},
            )
            image = visual_element_from_summary(
                "Site photo",
                element_type="image",
                source=source,
                document_id=document_id,
                path=path,
                page=3,
                confidence=0.61,
                labels=("Facade",),
                uncertainty="approximate from field notes",
                approximate=True,
                metadata={"figure_index": 2},
            )
            assert drawing is not None
            assert diagram is not None
            assert image is not None
            return [drawing, diagram, image]

        monkeypatch.setattr(ingestion_module.parsers, "parse_document", fake_parse_document)
        body_bytes = b"%PDF-1.7\nvisuals"

        response = client.post(
            "/api/ingest",
            files=[("files", ("visuals.pdf", io.BytesIO(body_bytes), "application/pdf"))],
        )

        assert response.status_code == 200, response.text
        body = response.json()
        assert body["ingested_files"] == 1
        assert body["ingested_chunks"] == 3
        records = fake_kb.dump()
        assert len(records) == 3

        registry_record = client.app.state.app_state.registry.get_by_hash(
            hashlib.sha256(body_bytes).hexdigest()
        )
        assert registry_record is not None
        assert registry_record.status == "indexed"
        assert registry_record.error is None
        assert registry_record.memory_ids == body["memory_ids"]
        assert body["memory_ids"] == [record["id"] for record in records]

        assert records[0]["content"] == (
            "[source=visuals.pdf; page=1; element=drawing; extraction=visual_summary; "
            "confidence=0.88]\n"
            "North stair sketch\n"
            "Labels: North stair; Access path\n"
            "Relationships: North stair -> Access path"
        )
        assert records[1]["content"] == (
            "[source=visuals.pdf; page=2; element=diagram; extraction=visual_summary; "
            "confidence=0.74]\n"
            "Workflow diagram\n"
            "Labels: Parser; Memory KB\n"
            "Relationships: Parser -> Ingest; Ingest -> KB"
        )
        assert records[2]["content"] == (
            "[source=visuals.pdf; page=3; element=image; extraction=visual_summary; "
            f"confidence=0.61; warnings={APPROXIMATE_VALUE_WARNING}]\n"
            "Site photo\n"
            "Labels: Facade\n"
            "Uncertainty: approximate from field notes"
        )

        assert [record["metadata"]["page"] for record in records] == [1, 2, 3]
        assert [record["metadata"]["element_type"] for record in records] == [
            "drawing",
            "diagram",
            "image",
        ]
        assert [record["metadata"]["extraction_mode"] for record in records] == [
            "visual_summary",
            "visual_summary",
            "visual_summary",
        ]
        assert [record["metadata"]["document_id"] for record in records] == [
            registry_record.document_id,
            registry_record.document_id,
            registry_record.document_id,
        ]
        assert [record["metadata"]["figure_index"] for record in records] == [0, 1, 2]

        drawing_metadata = records[0]["metadata"]
        assert drawing_metadata["path"] == registry_record.stored_path
        assert drawing_metadata["confidence"] == 0.88
        assert drawing_metadata["warnings"] == []
        assert drawing_metadata["visual_summary_chars"] == len("North stair sketch")
        assert drawing_metadata["labels"] == ["North stair", "Access path"]
        assert drawing_metadata["relationships"] == ["North stair -> Access path"]

        diagram_metadata = records[1]["metadata"]
        assert diagram_metadata["path"] == registry_record.stored_path
        assert diagram_metadata["confidence"] == 0.74
        assert diagram_metadata["warnings"] == []
        assert diagram_metadata["visual_summary_chars"] == len("Workflow diagram")
        assert diagram_metadata["labels"] == ["Parser", "Memory KB"]
        assert diagram_metadata["relationships"] == [
            "Parser -> Ingest",
            "Ingest -> KB",
        ]

        image_metadata = records[2]["metadata"]
        assert image_metadata["path"] == registry_record.stored_path
        assert image_metadata["confidence"] == 0.61
        assert image_metadata["warnings"] == [APPROXIMATE_VALUE_WARNING]
        assert image_metadata["visual_summary_chars"] == len("Site photo")
        assert image_metadata["labels"] == ["Facade"]
        assert image_metadata["approximate"] is True
        assert image_metadata["uncertainty"] == "approximate from field notes"

    def test_uploads_visual_elements_are_enriched_through_live_route(
        self,
        client: TestClient,
        fake_kb: FakeKB,
        tmp_path: Path,
        monkeypatch: pytest.MonkeyPatch,
    ) -> None:
        client.app.state.app_state.settings.documents_dir = str(tmp_path)
        analyzer, analysis_client = _build_visual_analyzer(
            responses={
                "drawing": VisualEnrichmentOutput(
                    summary="North stair sketch refined",
                    labels=["North stair", "Access path"],
                    relationships=["North stair -> Access path"],
                    confidence=0.96,
                ),
                "diagram": VisualEnrichmentOutput(
                    summary="Workflow diagram refined",
                    labels=["Parser", "Memory KB"],
                    relationships=["Parser -> Ingest", "Ingest -> KB"],
                    confidence=0.86,
                ),
                "image": VisualEnrichmentOutput(
                    summary="Site photo annotated",
                    labels=["Facade"],
                    uncertainty="approximate from field notes",
                    approximate=True,
                    confidence=0.79,
                ),
            },
        )
        client.app.state.app_state.document_analyzer = analyzer
        assert client.app.state.app_state.document_analyzer is analyzer

        parsed_elements: list[DocumentElement] = []

        def fake_parse_document(
            path: str,
            *,
            source: str,
            document_id: str | None = None,
        ) -> list[DocumentElement]:
            assert Path(path).exists()
            assert source == "visuals.pdf"
            assert document_id is not None
            paragraph = DocumentElement(
                document_id=document_id,
                source=source,
                path=path,
                page=1,
                element_type="paragraph",
                extraction_mode="pdf_text",
                content="cover page body",
                metadata={"section": "cover"},
            )
            drawing = visual_element_from_summary(
                "North stair sketch",
                element_type="drawing",
                source=source,
                document_id=document_id,
                path=path,
                page=2,
                confidence=0.88,
                metadata={"figure_index": 0},
            )
            diagram = visual_element_from_summary(
                "Workflow diagram",
                element_type="diagram",
                source=source,
                document_id=document_id,
                path=path,
                page=3,
                confidence=0.74,
                metadata={"figure_index": 1},
            )
            image = visual_element_from_summary(
                "Site photo",
                element_type="image",
                source=source,
                document_id=document_id,
                path=path,
                page=4,
                confidence=0.61,
                metadata={"figure_index": 2},
            )
            assert drawing is not None
            assert diagram is not None
            assert image is not None
            parsed_elements.extend([paragraph, drawing, diagram, image])
            return parsed_elements

        monkeypatch.setattr(ingestion_module.parsers, "parse_document", fake_parse_document)
        body_bytes = b"%PDF-1.7\nvisuals"

        response = client.post(
            "/api/ingest",
            files=[("files", ("visuals.pdf", io.BytesIO(body_bytes), "application/pdf"))],
        )

        assert response.status_code == 200, response.text
        body = response.json()
        assert body["ingested_files"] == 1
        assert body["ingested_chunks"] == 4
        records = fake_kb.dump()
        assert len(records) == 4

        registry_record = client.app.state.app_state.registry.get_by_hash(
            hashlib.sha256(body_bytes).hexdigest()
        )
        assert registry_record is not None
        assert registry_record.status == "indexed"
        assert registry_record.error is None
        assert registry_record.memory_ids == body["memory_ids"]
        assert body["memory_ids"] == [record["id"] for record in records]
        assert analysis_client.calls == parsed_elements[1:]

        assert records[0]["content"] == (
            "[source=visuals.pdf; page=1; element=paragraph; extraction=pdf_text]\ncover page body"
        )
        assert records[0]["metadata"] == {
            "document_id": registry_record.document_id,
            "source": "visuals.pdf",
            "path": registry_record.stored_path,
            "page": 1,
            "element_type": "paragraph",
            "extraction_mode": "pdf_text",
            "warnings": [],
            "section": "cover",
        }

        assert records[1]["content"] == (
            "[source=visuals.pdf; page=2; element=drawing; extraction=visual_summary; "
            "confidence=0.96]\n"
            "North stair sketch refined\n"
            "Labels: North stair; Access path\n"
            "Relationships: North stair -> Access path"
        )
        assert records[1]["metadata"] == {
            "document_id": registry_record.document_id,
            "source": "visuals.pdf",
            "path": registry_record.stored_path,
            "page": 2,
            "element_type": "drawing",
            "extraction_mode": "visual_summary",
            "confidence": 0.96,
            "warnings": [],
            "visual_summary_chars": len("North stair sketch refined"),
            "labels": ["North stair", "Access path"],
            "relationships": ["North stair -> Access path"],
            "figure_index": 0,
            "analysis_provider": "openai",
            "analysis_model": "fake-visual-model",
            "analysis_mode": "visual_only",
            "analysis_status": "enriched",
            "analysis_source_element_type": "drawing",
            "analysis_source_extraction_mode": "visual_summary",
            "analysis_source_confidence": 0.88,
        }

        assert records[2]["content"] == (
            "[source=visuals.pdf; page=3; element=diagram; extraction=visual_summary; "
            "confidence=0.86]\n"
            "Workflow diagram refined\n"
            "Labels: Parser; Memory KB\n"
            "Relationships: Parser -> Ingest; Ingest -> KB"
        )
        assert records[2]["metadata"] == {
            "document_id": registry_record.document_id,
            "source": "visuals.pdf",
            "path": registry_record.stored_path,
            "page": 3,
            "element_type": "diagram",
            "extraction_mode": "visual_summary",
            "confidence": 0.86,
            "warnings": [],
            "visual_summary_chars": len("Workflow diagram refined"),
            "labels": ["Parser", "Memory KB"],
            "relationships": ["Parser -> Ingest", "Ingest -> KB"],
            "figure_index": 1,
            "analysis_provider": "openai",
            "analysis_model": "fake-visual-model",
            "analysis_mode": "visual_only",
            "analysis_status": "enriched",
            "analysis_source_element_type": "diagram",
            "analysis_source_extraction_mode": "visual_summary",
            "analysis_source_confidence": 0.74,
        }

        assert records[3]["content"] == (
            "[source=visuals.pdf; page=4; element=image; extraction=visual_summary; "
            "confidence=0.79; warnings=approximate_values]\n"
            "Site photo annotated\n"
            "Labels: Facade\n"
            "Uncertainty: approximate from field notes"
        )
        assert records[3]["metadata"] == {
            "document_id": registry_record.document_id,
            "source": "visuals.pdf",
            "path": registry_record.stored_path,
            "page": 4,
            "element_type": "image",
            "extraction_mode": "visual_summary",
            "confidence": 0.79,
            "warnings": [APPROXIMATE_VALUE_WARNING],
            "visual_summary_chars": len("Site photo annotated"),
            "labels": ["Facade"],
            "uncertainty": "approximate from field notes",
            "approximate": True,
            "figure_index": 2,
            "analysis_provider": "openai",
            "analysis_model": "fake-visual-model",
            "analysis_mode": "visual_only",
            "analysis_status": "enriched",
            "analysis_source_element_type": "image",
            "analysis_source_extraction_mode": "visual_summary",
            "analysis_source_confidence": 0.61,
        }

    async def test_uploads_mixed_typed_evidence_through_live_route_and_kb_recall(
        self,
        client: TestClient,
        fake_kb: FakeKB,
        tmp_path: Path,
        monkeypatch: pytest.MonkeyPatch,
    ) -> None:
        client.app.state.app_state.settings.documents_dir = str(tmp_path)
        analyzer, analysis_client = _build_visual_analyzer(
            responses={
                "image": VisualEnrichmentOutput(
                    summary="Site photo annotated",
                    labels=["Facade"],
                    relationships=["Facade -> Entry"],
                    uncertainty="approximate from field notes",
                    approximate=True,
                    confidence=0.79,
                ),
            },
        )
        client.app.state.app_state.document_analyzer = analyzer
        assert client.app.state.app_state.document_analyzer is analyzer
        kb = client.app.state.app_state.kb
        assert kb is fake_kb

        parsed_elements: list[DocumentElement] = []

        def fake_parse_document(
            path: str,
            *,
            source: str,
            document_id: str | None = None,
        ) -> list[DocumentElement]:
            assert Path(path).exists()
            assert Path(path).parent == tmp_path
            assert source == "mixed.pdf"
            assert document_id is not None
            paragraph = DocumentElement(
                document_id=document_id,
                source=source,
                path=path,
                page=1,
                element_type="paragraph",
                extraction_mode="pdf_text",
                content="Mixed upload keeps the route searchable.",
                metadata={"section": "intro"},
            )
            table = table_element_from_rows(
                [["Room", "Width", "Notes"], ["North stair", "42", "egress path"]],
                document_id=document_id,
                source=source,
                path=path,
                page=2,
                confidence=0.82,
                warnings=("merged_cells",),
                metadata={"table_index": 0},
            )
            ocr = ocr_element_from_text(
                "  Recovered\nsheet\tnote from scan  ",
                document_id=document_id,
                source=source,
                path=path,
                page=3,
                confidence=0.41,
                warnings=("low_text_page", "ocr_low_confidence"),
                low_text_threshold=20,
                metadata={"ocr_engine": "fake"},
            )
            visual = visual_element_from_summary(
                "Site photo",
                element_type="image",
                source=source,
                document_id=document_id,
                path=path,
                page=4,
                confidence=0.62,
                warnings=("field_review",),
                metadata={"figure_index": 7},
            )
            assert table is not None
            assert ocr is not None
            assert visual is not None
            parsed_elements.extend([paragraph, table, ocr, visual])
            return parsed_elements

        monkeypatch.setattr(ingestion_module.parsers, "parse_document", fake_parse_document)
        body_bytes = b"%PDF-1.7\nmixed typed evidence"

        response = client.post(
            "/api/ingest",
            files=[("files", ("mixed.pdf", io.BytesIO(body_bytes), "application/pdf"))],
        )

        assert response.status_code == 200, response.text
        body = response.json()
        assert body["ingested_files"] == 1
        assert body["ingested_chunks"] == 4
        records = fake_kb.dump()
        assert len(records) == 4

        registry_record = client.app.state.app_state.registry.get_by_hash(
            hashlib.sha256(body_bytes).hexdigest()
        )
        assert registry_record is not None
        assert registry_record.status == "indexed"
        assert registry_record.error is None
        assert registry_record.memory_ids == body["memory_ids"]
        assert body["memory_ids"] == [record["id"] for record in records]
        assert Path(registry_record.stored_path).parent == tmp_path
        assert Path(registry_record.stored_path).exists()
        assert analysis_client.calls == [parsed_elements[3]]

        document_id = registry_record.document_id
        stored_path = registry_record.stored_path
        paragraph_content = (
            "[source=mixed.pdf; page=1; element=paragraph; extraction=pdf_text]\n"
            "Mixed upload keeps the route searchable."
        )
        table_content = (
            "[source=mixed.pdf; page=2; element=table; extraction=pdf_table; "
            "confidence=0.82; warnings=merged_cells]\n"
            "| Room | Width | Notes |\n"
            "| --- | --- | --- |\n"
            "| North stair | 42 | egress path |"
        )
        ocr_content = (
            "[source=mixed.pdf; page=3; element=ocr_text; extraction=ocr; "
            "confidence=0.41; warnings=low_text_page,ocr_low_confidence]\n"
            "Recovered sheet note from scan"
        )
        visual_content = (
            f"[source=mixed.pdf; page=4; element=image; extraction=visual_summary; "
            f"confidence=0.79; warnings=field_review,{APPROXIMATE_VALUE_WARNING}]\n"
            "Site photo annotated\n"
            "Labels: Facade\n"
            "Relationships: Facade -> Entry\n"
            "Uncertainty: approximate from field notes"
        )

        assert records[0]["content"] == paragraph_content
        assert records[0]["metadata"] == {
            "document_id": document_id,
            "source": "mixed.pdf",
            "path": stored_path,
            "page": 1,
            "element_type": "paragraph",
            "extraction_mode": "pdf_text",
            "warnings": [],
            "section": "intro",
        }
        assert records[1]["content"] == table_content
        assert records[1]["metadata"] == {
            "document_id": document_id,
            "source": "mixed.pdf",
            "path": stored_path,
            "page": 2,
            "element_type": "table",
            "extraction_mode": "pdf_table",
            "confidence": 0.82,
            "warnings": ["merged_cells"],
            "table_rows": 2,
            "table_columns": 3,
            "table_index": 0,
        }
        assert records[2]["content"] == ocr_content
        assert records[2]["metadata"] == {
            "document_id": document_id,
            "source": "mixed.pdf",
            "path": stored_path,
            "page": 3,
            "element_type": "ocr_text",
            "extraction_mode": "ocr",
            "confidence": 0.41,
            "warnings": ["low_text_page", "ocr_low_confidence"],
            "ocr_text_chars": len("Recovered sheet note from scan"),
            "low_text_threshold": 20,
            "ocr_engine": "fake",
        }
        assert records[3]["content"] == visual_content
        assert records[3]["metadata"] == {
            "document_id": document_id,
            "source": "mixed.pdf",
            "path": stored_path,
            "page": 4,
            "element_type": "image",
            "extraction_mode": "visual_summary",
            "confidence": 0.79,
            "warnings": ["field_review", APPROXIMATE_VALUE_WARNING],
            "visual_summary_chars": len("Site photo annotated"),
            "labels": ["Facade"],
            "relationships": ["Facade -> Entry"],
            "uncertainty": "approximate from field notes",
            "approximate": True,
            "figure_index": 7,
            "analysis_provider": "openai",
            "analysis_model": "fake-visual-model",
            "analysis_mode": "visual_only",
            "analysis_status": "enriched",
            "analysis_source_element_type": "image",
            "analysis_source_extraction_mode": "visual_summary",
            "analysis_source_confidence": 0.62,
        }

        kb_recall = next(tool for tool in build_kb_tools(kb) if tool.name == "kb_recall")
        for query, record in (
            ("route searchable", records[0]),
            ("egress path", records[1]),
            ("recovered sheet note", records[2]),
            ("field notes", records[3]),
        ):
            result = await kb_recall.ainvoke({"query": query, "k": 5})
            assert result == f"- ({record['id']}) {record['content']}"

    @pytest.mark.parametrize(
        ("failure", "expected_warning", "expected_status"),
        [
            pytest.param(
                RuntimeError("boom"),
                "openai_enrichment_failed",
                "failed",
                id="exception",
            ),
            pytest.param(
                DocumentAnalysisRefusalError("policy refusal"),
                "openai_enrichment_refused",
                "refused",
                id="refusal",
            ),
        ],
    )
    def test_uploads_visual_analysis_failures_are_indexed_with_warning(
        self,
        client: TestClient,
        fake_kb: FakeKB,
        tmp_path: Path,
        monkeypatch: pytest.MonkeyPatch,
        failure: Exception,
        expected_warning: str,
        expected_status: str,
    ) -> None:
        client.app.state.app_state.settings.documents_dir = str(tmp_path)
        analyzer, analysis_client = _build_visual_analyzer(exception=failure)
        client.app.state.app_state.document_analyzer = analyzer
        assert client.app.state.app_state.document_analyzer is analyzer

        parsed_elements: list[DocumentElement] = []

        def fake_parse_document(
            path: str,
            *,
            source: str,
            document_id: str | None = None,
        ) -> list[DocumentElement]:
            assert Path(path).exists()
            assert source == "visuals.pdf"
            assert document_id is not None
            chart = visual_element_from_summary(
                "Original visual summary",
                element_type="chart",
                source=source,
                document_id=document_id,
                path=path,
                page=4,
                confidence=0.83,
                metadata={"figure_index": 8},
            )
            assert chart is not None
            parsed_elements.append(chart)
            return [chart]

        monkeypatch.setattr(ingestion_module.parsers, "parse_document", fake_parse_document)
        body_bytes = b"%PDF-1.7\nvisuals"

        response = client.post(
            "/api/ingest",
            files=[("files", ("visuals.pdf", io.BytesIO(body_bytes), "application/pdf"))],
        )

        assert response.status_code == 200, response.text
        body = response.json()
        assert body["ingested_files"] == 1
        assert body["ingested_chunks"] == 1
        records = fake_kb.dump()
        assert len(records) == 1

        registry_record = client.app.state.app_state.registry.get_by_hash(
            hashlib.sha256(body_bytes).hexdigest()
        )
        assert registry_record is not None
        assert registry_record.status == "indexed"
        assert registry_record.error is None
        assert registry_record.memory_ids == body["memory_ids"]
        assert body["memory_ids"] == [record["id"] for record in records]
        assert analysis_client.calls == parsed_elements

        assert records[0]["content"] == (
            f"[source=visuals.pdf; page=4; element=chart; extraction=visual_summary; "
            f"confidence=0.83; warnings={expected_warning}]\n"
            "Original visual summary"
        )
        assert records[0]["metadata"] == {
            "document_id": registry_record.document_id,
            "source": "visuals.pdf",
            "path": registry_record.stored_path,
            "page": 4,
            "element_type": "chart",
            "extraction_mode": "visual_summary",
            "confidence": 0.83,
            "warnings": [expected_warning],
            "visual_summary_chars": len("Original visual summary"),
            "figure_index": 8,
            "analysis_provider": "openai",
            "analysis_model": "fake-visual-model",
            "analysis_mode": "visual_only",
            "analysis_status": expected_status,
            "analysis_source_element_type": "chart",
            "analysis_source_extraction_mode": "visual_summary",
            "analysis_source_confidence": 0.83,
        }

    def test_uploads_pdf_ocr_element_as_standalone_memory_with_warning(
        self,
        client: TestClient,
        fake_kb: FakeKB,
        tmp_path: Path,
        monkeypatch: pytest.MonkeyPatch,
    ) -> None:
        client.app.state.app_state.settings.documents_dir = str(tmp_path)

        def fake_parse_document(
            path: str,
            *,
            source: str,
            document_id: str | None = None,
        ) -> list[DocumentElement]:
            assert Path(path).exists()
            assert source == "scan.pdf"
            assert document_id is not None
            ocr_element = ocr_element_from_text(
                "  recovered\n sheet\tnote  ",
                document_id=document_id,
                source=source,
                path=path,
                page=2,
                confidence=0.41,
                warnings=("low_text_page", "ocr_low_confidence"),
                low_text_threshold=20,
                metadata={"ocr_engine": "fake"},
            )
            assert ocr_element is not None
            return [
                DocumentElement(
                    document_id=document_id,
                    source=source,
                    path=path,
                    page=1,
                    element_type="paragraph",
                    extraction_mode="pdf_text",
                    content="parsed cover sheet",
                ),
                ocr_element,
            ]

        monkeypatch.setattr(ingestion_module.parsers, "parse_document", fake_parse_document)
        body_bytes = b"%PDF-1.7\nocr"

        response = client.post(
            "/api/ingest",
            files=[("files", ("scan.pdf", io.BytesIO(body_bytes), "application/pdf"))],
        )

        assert response.status_code == 200, response.text
        body = response.json()
        assert body["ingested_files"] == 1
        assert body["ingested_chunks"] == 2
        records = fake_kb.dump()
        assert len(records) == 2

        registry_record = client.app.state.app_state.registry.get_by_hash(
            hashlib.sha256(body_bytes).hexdigest()
        )
        assert registry_record is not None
        assert registry_record.status == "indexed"
        assert registry_record.error is None
        assert registry_record.memory_ids == body["memory_ids"]
        assert body["memory_ids"] == [record["id"] for record in records]

        assert records[0]["content"].startswith(
            "[source=scan.pdf; page=1; element=paragraph; extraction=pdf_text]\n"
        )
        assert records[0]["content"].endswith("parsed cover sheet")

        content = records[1]["content"]
        assert content.startswith(
            "[source=scan.pdf; page=2; element=ocr_text; extraction=ocr; "
            "confidence=0.41; warnings=low_text_page,ocr_low_confidence]\n"
        )
        assert content.endswith("recovered sheet note")

        metadata = records[1]["metadata"]
        assert metadata["document_id"] == registry_record.document_id
        assert metadata["source"] == "scan.pdf"
        assert metadata["path"] == registry_record.stored_path
        assert metadata["page"] == 2
        assert metadata["element_type"] == "ocr_text"
        assert metadata["extraction_mode"] == "ocr"
        assert metadata["confidence"] == 0.41
        assert metadata["warnings"] == ["low_text_page", "ocr_low_confidence"]
        assert metadata["ocr_text_chars"] == len("recovered sheet note")
        assert metadata["low_text_threshold"] == 20
        assert metadata["ocr_engine"] == "fake"

    def test_uploads_pdf_table_element_as_standalone_memory_with_warning(
        self,
        client: TestClient,
        fake_kb: FakeKB,
        tmp_path: Path,
        monkeypatch: pytest.MonkeyPatch,
    ) -> None:
        client.app.state.app_state.settings.documents_dir = str(tmp_path)

        def fake_parse_document(
            path: str,
            *,
            source: str,
            document_id: str | None = None,
        ) -> list[DocumentElement]:
            assert Path(path).exists()
            assert source == "tables.pdf"
            assert document_id is not None
            element = table_element_from_rows(
                [["Room", "Area", "Notes"], ["A101", "42 m2"]],
                document_id=document_id,
                source=source,
                path=path,
                page=5,
                confidence=0.78,
                warnings=("merged_cells",),
                metadata={"table_index": 0},
            )
            assert element is not None
            return [element]

        monkeypatch.setattr(ingestion_module.parsers, "parse_document", fake_parse_document)
        body_bytes = b"%PDF-1.7\ntable"

        response = client.post(
            "/api/ingest",
            files=[("files", ("tables.pdf", io.BytesIO(body_bytes), "application/pdf"))],
        )

        assert response.status_code == 200, response.text
        body = response.json()
        assert body["ingested_files"] == 1
        assert body["ingested_chunks"] == 1
        records = fake_kb.dump()
        assert len(records) == 1

        registry_record = client.app.state.app_state.registry.get_by_hash(
            hashlib.sha256(body_bytes).hexdigest()
        )
        assert registry_record is not None
        assert registry_record.status == "indexed"
        assert registry_record.error is None
        assert registry_record.memory_ids == body["memory_ids"]
        assert body["memory_ids"] == [records[0]["id"]]

        content = records[0]["content"]
        assert content.startswith(
            "[source=tables.pdf; page=5; element=table; extraction=pdf_table; "
            f"confidence=0.78; warnings=merged_cells,{RAGGED_TABLE_WARNING}]\n"
        )
        assert content.endswith("| Room | Area | Notes |\n| --- | --- | --- |\n| A101 | 42 m2 |  |")
        metadata = records[0]["metadata"]
        assert metadata["document_id"] == registry_record.document_id
        assert metadata["source"] == "tables.pdf"
        assert metadata["path"] == registry_record.stored_path
        assert metadata["page"] == 5
        assert metadata["element_type"] == "table"
        assert metadata["extraction_mode"] == "pdf_table"
        assert metadata["confidence"] == 0.78
        assert metadata["warnings"] == ["merged_cells", RAGGED_TABLE_WARNING]
        assert metadata["table_rows"] == 2
        assert metadata["table_columns"] == 3
        assert metadata["table_index"] == 0

    def test_uploads_docx_emits_summary_heading_paragraph_and_table_memories(
        self,
        client: TestClient,
        fake_kb: FakeKB,
        tmp_path: Path,
    ) -> None:
        client.app.state.app_state.settings.documents_dir = str(tmp_path)

        document = docx.Document()
        document.core_properties.title = "Loads Spec"
        document.core_properties.author = "Test Engineer"
        document.add_heading("Loads", level=1)
        document.add_paragraph("Live load: 5 kN per square meter.")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Spec"
        table.cell(0, 1).text = "Value"
        table.cell(1, 0).text = "Height"
        table.cell(1, 1).text = "10 m"

        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        body_bytes = buffer.getvalue()

        response = client.post(
            "/api/ingest",
            files=[
                (
                    "files",
                    (
                        "loads.docx",
                        io.BytesIO(body_bytes),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                )
            ],
        )

        assert response.status_code == 200, response.text
        body = response.json()
        assert body["ingested_files"] == 1
        assert len(body["memory_ids"]) >= 4

        records = fake_kb.dump()
        assert body["memory_ids"] == [record["id"] for record in records]

        registry_record = client.app.state.app_state.registry.get_by_hash(
            hashlib.sha256(body_bytes).hexdigest()
        )
        assert registry_record is not None
        assert registry_record.status == "indexed"
        assert registry_record.error is None
        assert registry_record.memory_ids == body["memory_ids"]
        assert Path(registry_record.stored_path).exists()

        extraction_modes = {record["metadata"]["extraction_mode"] for record in records}
        assert {
            "docx_summary",
            "docx_heading",
            "docx_paragraph",
            "docx_table",
        } <= extraction_modes

        heading_record = next(
            record for record in records if record["metadata"]["extraction_mode"] == "docx_heading"
        )
        assert heading_record["metadata"].get("section_heading") is None
        assert heading_record["content"].startswith("[source=loads.docx;")
        assert "element=heading; extraction=docx_heading" in heading_record["content"]

        paragraph_record = next(
            record
            for record in records
            if record["metadata"]["extraction_mode"] == "docx_paragraph"
        )
        assert paragraph_record["metadata"]["section_heading"] == "Loads"
        assert paragraph_record["content"].startswith("[source=loads.docx;")
        assert "element=paragraph; extraction=docx_paragraph" in paragraph_record["content"]

        table_record = next(
            record for record in records if record["metadata"]["extraction_mode"] == "docx_table"
        )
        assert table_record["metadata"]["table_rows"] == 2
        assert table_record["metadata"]["table_columns"] == 2
        assert table_record["metadata"]["section_heading"] == "Loads"
        assert "| Spec | Value |" in table_record["content"]
        assert "| Height | 10 m |" in table_record["content"]

        summary_record = next(
            record for record in records if record["metadata"]["extraction_mode"] == "docx_summary"
        )
        assert summary_record["metadata"]["subject"] == "engineering_narrative"
        assert summary_record["metadata"]["paragraph_count"] == 2

    def test_uploads_pdf_skips_empty_page_elements_and_preserves_page_gap(
        self,
        client: TestClient,
        fake_kb: FakeKB,
        tmp_path: Path,
        monkeypatch: pytest.MonkeyPatch,
    ) -> None:
        client.app.state.app_state.settings.documents_dir = str(tmp_path)

        def fake_parse_document(
            path: str,
            *,
            source: str,
            document_id: str | None = None,
        ) -> list[DocumentElement]:
            return [
                DocumentElement(
                    document_id=document_id,
                    source=source,
                    path=path,
                    page=1,
                    element_type="paragraph",
                    extraction_mode="pdf_text",
                    content="first page body",
                ),
                DocumentElement(
                    document_id=document_id,
                    source=source,
                    path=path,
                    page=2,
                    element_type="paragraph",
                    extraction_mode="pdf_text",
                    content="   ",
                ),
                DocumentElement(
                    document_id=document_id,
                    source=source,
                    path=path,
                    page=3,
                    element_type="paragraph",
                    extraction_mode="pdf_text",
                    content="third page body",
                ),
            ]

        monkeypatch.setattr(ingestion_module.parsers, "parse_document", fake_parse_document)

        response = client.post(
            "/api/ingest",
            files=[("files", ("gapped.pdf", io.BytesIO(b"%PDF-1.7\ngapped"), "application/pdf"))],
        )

        assert response.status_code == 200, response.text
        body = response.json()
        assert body["ingested_files"] == 1
        assert body["ingested_chunks"] == 2
        records = fake_kb.dump()
        assert len(records) == 2
        assert [record["metadata"]["page"] for record in records] == [1, 3]
        assert records[0]["content"].startswith("[source=gapped.pdf; page=1;")
        assert records[1]["content"].startswith("[source=gapped.pdf; page=3;")

    def test_uploads_multiple_files(self, client: TestClient, tmp_path: Path) -> None:
        client.app.state.app_state.settings.documents_dir = str(tmp_path)
        files = [
            ("files", ("a.txt", io.BytesIO(b"alpha"), "text/plain")),
            ("files", ("b.md", io.BytesIO(b"# beta"), "text/markdown")),
        ]
        r = client.post("/api/ingest", files=files)
        assert r.status_code == 200
        assert r.json()["ingested_files"] == 2
        assert len(list(tmp_path.iterdir())) == 2
