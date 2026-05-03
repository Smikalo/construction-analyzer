"""Tests for document parser adapters and extension dispatch."""

from __future__ import annotations

from pathlib import Path

import docx
import pytest

import app.services.parsers as parsers


class FakePage:
    def __init__(self, text: str | None) -> None:
        self._text = text

    def extract_text(self) -> str | None:
        return self._text


def test_parse_markdown_returns_single_markdown_element(tmp_path: Path) -> None:
    path = tmp_path / "scope.md"
    path.write_text("# Scope\n\nBuild the wall.", encoding="utf-8")

    elements = parsers.parse_document(str(path), source="scope.md", document_id="doc-md")

    assert len(elements) == 1
    element = elements[0]
    assert element.document_id == "doc-md"
    assert element.source == "scope.md"
    assert element.path == str(path)
    assert element.page is None
    assert element.element_type == "paragraph"
    assert element.extraction_mode == "markdown"
    assert element.content == "# Scope\n\nBuild the wall."
    assert element.confidence is None
    assert element.warnings == ()
    assert element.metadata == {}


def test_parse_text_returns_single_text_element(tmp_path: Path) -> None:
    path = tmp_path / "notes.txt"
    path.write_text("plain note body", encoding="utf-8")

    elements = parsers.parse_document(str(path), source="notes.txt", document_id="doc-txt")

    assert len(elements) == 1
    element = elements[0]
    assert element.document_id == "doc-txt"
    assert element.source == "notes.txt"
    assert element.path == str(path)
    assert element.page is None
    assert element.element_type == "paragraph"
    assert element.extraction_mode == "text"
    assert element.content == "plain note body"


def test_parse_pdf_emits_non_empty_pages_with_original_page_numbers(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    path = tmp_path / "multi.pdf"
    pages = [FakePage("page one body"), FakePage(""), FakePage("page three body")]

    class FakePdfReader:
        def __init__(self, reader_path: str) -> None:
            assert reader_path == str(path)
            self.pages = pages

    monkeypatch.setattr(parsers, "PdfReader", FakePdfReader)

    elements = parsers.parse_document(str(path), source="multi.pdf", document_id="doc-pdf")

    assert [element.page for element in elements] == [1, 3]
    assert [element.content for element in elements] == ["page one body", "page three body"]
    assert {element.extraction_mode for element in elements} == {"pdf_text"}
    assert {element.element_type for element in elements} == {"paragraph"}
    assert {element.document_id for element in elements} == {"doc-pdf"}
    assert {element.source for element in elements} == {"multi.pdf"}
    assert {element.path for element in elements} == {str(path)}
    assert all(element.confidence is None for element in elements)
    assert all(element.warnings == () for element in elements)


def test_parse_document_dispatches_docx_to_extract_docx(tmp_path: Path) -> None:
    path = tmp_path / "dispatch.docx"
    document = docx.Document()
    document.add_heading("Dispatch Heading", level=1)
    document.add_paragraph("Dispatch paragraph")
    document.save(path)

    elements = parsers.parse_document(
        str(path),
        source="dispatch.docx",
        document_id="dispatch-1",
    )

    assert len(elements) >= 3
    assert {element.document_id for element in elements} == {"dispatch-1"}
    assert {element.source for element in elements} == {"dispatch.docx"}
    assert {element.path for element in elements} == {str(path)}
    assert any(element.extraction_mode == "docx_summary" for element in elements)
    assert any(element.extraction_mode == "docx_heading" for element in elements)
    assert any(element.extraction_mode == "docx_paragraph" for element in elements)


def test_empty_pdf_returns_no_elements(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    path = tmp_path / "empty.pdf"

    class FakePdfReader:
        def __init__(self, reader_path: str) -> None:
            assert reader_path == str(path)
            self.pages: list[FakePage] = []

    monkeypatch.setattr(parsers, "PdfReader", FakePdfReader)

    assert parsers.parse_document(str(path), source="empty.pdf") == []


@pytest.mark.parametrize("filename", ["empty.md", "empty.markdown", "empty.txt"])
def test_empty_markdown_and_text_files_return_no_elements(tmp_path: Path, filename: str) -> None:
    path = tmp_path / filename
    path.write_text("  \n\t  ", encoding="utf-8")

    assert parsers.parse_document(str(path), source=filename) == []


def test_parse_document_returns_empty_for_unsupported_extension(tmp_path: Path) -> None:
    path = tmp_path / "image.png"
    path.write_bytes(b"not parsed")

    assert parsers.parse_document(str(path), source="image.png", document_id="doc-img") == []
