"""Tests for DOCX extraction helpers."""

from __future__ import annotations

from pathlib import Path

import docx

from app.services.docx_elements import extract_docx


def _save_docx(document, tmp_path: Path, filename: str) -> Path:
    path = tmp_path / filename
    document.save(path)
    return path


def test_extract_docx_emits_file_summary_with_core_properties_and_paragraph_count(
    tmp_path: Path,
) -> None:
    document = docx.Document()
    document.core_properties.title = "Loads Spec"
    document.core_properties.author = "Test Engineer"
    document.add_paragraph("First paragraph")
    document.add_paragraph("Second paragraph")
    document.add_paragraph("Third paragraph")

    path = _save_docx(document, tmp_path, "sample.docx")

    elements = extract_docx(str(path), source="sample.docx", document_id="doc-1")

    assert len(elements) == 4
    element = elements[0]
    assert element.document_id == "doc-1"
    assert element.source == "sample.docx"
    assert element.path == str(path)
    assert element.element_type == "file_summary"
    assert element.extraction_mode == "docx_summary"
    assert "Loads Spec" in element.content
    assert "Paragraphs: 3" in element.content
    assert element.metadata["subject"] == "engineering_narrative"
    assert element.metadata["paragraph_count"] == 3
    assert element.metadata["docx_title"] == "Loads Spec"
    assert element.metadata["docx_author"] == "Test Engineer"


def test_extract_docx_tracks_heading_sections_across_paragraphs(
    tmp_path: Path,
) -> None:
    document = docx.Document()
    document.add_heading("Site Prep", level=1)
    document.add_paragraph("Excavation begins")

    path = _save_docx(document, tmp_path, "heading_sections.docx")

    elements = extract_docx(str(path), source="heading_sections.docx", document_id="doc-2")

    assert [element.element_type for element in elements] == [
        "file_summary",
        "heading",
        "paragraph",
    ]

    heading = elements[1]
    paragraph = elements[2]

    assert heading.extraction_mode == "docx_heading"
    assert heading.content == "Site Prep"
    assert heading.metadata["block_index"] == 0
    assert heading.metadata["style_name"] == "Heading 1"
    assert "section_heading" not in heading.metadata

    assert paragraph.extraction_mode == "docx_paragraph"
    assert paragraph.content == "Excavation begins"
    assert paragraph.metadata["block_index"] == 1
    assert paragraph.metadata["style_name"] == "Normal"
    assert paragraph.metadata["section_heading"] == "Site Prep"


def test_extract_docx_renders_tables_with_dimensions_and_docx_mode(
    tmp_path: Path,
) -> None:
    document = docx.Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "C"
    table.cell(1, 1).text = "D"

    path = _save_docx(document, tmp_path, "table.docx")

    elements = extract_docx(str(path), source="table.docx", document_id="doc-3")

    assert len(elements) == 2
    table_element = elements[1]
    assert table_element.element_type == "table"
    assert table_element.extraction_mode == "docx_table"
    assert table_element.metadata["block_index"] == 0
    assert table_element.metadata["subject"] == "engineering_narrative"
    assert table_element.metadata["table_rows"] == 2
    assert table_element.metadata["table_columns"] == 2


def test_extract_docx_preserves_document_order_and_block_indices(
    tmp_path: Path,
) -> None:
    document = docx.Document()
    document.add_heading("Foundation", level=1)
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Pile cap"
    document.add_paragraph("Backfill starts")

    path = _save_docx(document, tmp_path, "order.docx")

    elements = extract_docx(str(path), source="order.docx", document_id="doc-4")
    body_elements = elements[1:]

    assert [element.element_type for element in body_elements] == [
        "heading",
        "table",
        "paragraph",
    ]
    assert [element.metadata["block_index"] for element in body_elements] == [0, 1, 2]
    assert body_elements[0].metadata.get("section_heading") is None
    assert body_elements[1].metadata["section_heading"] == "Foundation"
    assert body_elements[2].metadata["section_heading"] == "Foundation"


def test_extract_docx_skips_empty_paragraphs_without_consuming_block_indices(
    tmp_path: Path,
) -> None:
    document = docx.Document()
    document.add_heading("Sequence", level=1)
    document.add_paragraph("   ")
    document.add_paragraph("Kept paragraph")

    path = _save_docx(document, tmp_path, "empty_paragraph.docx")

    elements = extract_docx(str(path), source="empty_paragraph.docx", document_id="doc-5")
    body_elements = elements[1:]

    assert [element.element_type for element in body_elements] == ["heading", "paragraph"]
    assert [element.metadata["block_index"] for element in body_elements] == [0, 1]
    assert body_elements[1].content == "Kept paragraph"
    assert body_elements[1].metadata["section_heading"] == "Sequence"


def test_extract_docx_flags_nested_tables_with_unsupported_structure_warning(
    tmp_path: Path,
) -> None:
    document = docx.Document()
    outer_table = document.add_table(rows=1, cols=1)
    outer_cell = outer_table.cell(0, 0)
    outer_cell.text = "Outer"
    nested_table = outer_cell.add_table(rows=1, cols=1)
    nested_table.cell(0, 0).text = "Inner"

    path = _save_docx(document, tmp_path, "nested_table.docx")

    elements = extract_docx(str(path), source="nested_table.docx", document_id="doc-6")

    assert len(elements) == 2
    table_element = elements[1]
    assert table_element.element_type == "table"
    assert table_element.extraction_mode == "docx_table"
    assert table_element.warnings == ("unsupported_structure",)
    assert table_element.metadata["block_index"] == 0
    assert table_element.metadata["table_rows"] == 1
    assert table_element.metadata["table_columns"] == 1
